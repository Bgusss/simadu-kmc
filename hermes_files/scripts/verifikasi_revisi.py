from docx import Document
d = Document(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_RENDIANSYAH.docx")
print("=== VERIFIKASI PARAGRAF AKHIR ===")
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if '3.2.3' in t or '3.2.4' in t or '3.2.5' in t or '3.2.6' in t or '3.2.7' in t:
        print(f"P{i}: {t}")
print("=== DETEKSI DUPLIKAT DAN HEADING GANDA ===")
headings = [p.text.strip() for p in d.paragraphs if p.text.strip().startswith(('A.', 'B.', 'C.', 'D.', 'E.')) or 'Perancangan UML' in p.text]
print("Headings in UML Section:", headings)

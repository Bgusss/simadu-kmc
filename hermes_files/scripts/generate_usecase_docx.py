import os
from docx import Document

def main():
    doc_path = r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\Usecase_scenario_fix2.docx"
    
    # Create a fresh new document
    doc = Document()
    
    # Set standard academic formatting styles (Times New Roman, Font sizes, etc. can be modified in Word)
    # The default margins and font properties will inherit Word defaults.
    
    doc.add_heading('SKENARIO USE CASE SISTEM', level=1)
    doc.add_paragraph('Berikut adalah detail skenario use case untuk setiap fungsi yang terdapat pada Use Case Diagram.')
    
    scenarios = [
        {
            "id": "SUC-01",
            "name": "Login",
            "desc": "Aktor melakukan autentikasi untuk masuk ke dalam sistem dengan memasukkan email dan kata sandi agar dapat mengakses fitur sesuai hak aksesnya.",
            "actor": "Admin KMC, Pengguna OPD",
            "pre": "Aktor belum masuk ke dalam sistem dan berada di halaman login.",
            "post": "Aktor berhasil masuk dan diarahkan ke dashboard masing-masing.",
            "normal": [
                "Aktor mengakses halaman login.",
                "Aktor memasukkan email/username dan kata sandi.",
                "Aktor menekan tombol masuk.",
                "Sistem memvalidasi kredensial aktor.",
                "Sistem mengarahkan Admin KMC ke Dashboard Admin atau Pengguna OPD ke Dashboard OPD."
            ],
            "alt": [
                "Aktor memasukkan kredensial yang tidak valid.",
                "Sistem menampilkan pesan peringatan bahwa email/username atau kata sandi salah.",
                "Sistem mengembalikan aktor ke halaman login."
            ]
        },
        {
            "id": "SUC-02",
            "name": "Melihat Dashboard Admin",
            "desc": "Admin KMC memantau ringkasan informasi, statistik singkat, dan notifikasi aduan KMC secara real-time.",
            "actor": "Admin KMC",
            "pre": "Admin KMC telah berhasil masuk ke sistem.",
            "post": "Admin KMC melihat ringkasan data aduan dan tiket terkini.",
            "normal": [
                "Admin KMC memilih menu Dashboard.",
                "Sistem menghitung dan memuat ringkasan data tiket, notifikasi terbaru, dan notifikasi prioritas tinggi.",
                "Sistem menampilkan informasi tersebut pada halaman dashboard Admin."
            ],
            "alt": [
                "Tidak ada."
            ]
        },
        {
            "id": "SUC-03",
            "name": "Memantau Notifikasi dan Hasil Klasifikasi",
            "desc": "Admin KMC melihat daftar pesan masuk dari media sosial yang telah disaring dan diklasifikasikan oleh layanan AI.",
            "actor": "Admin KMC",
            "pre": "Admin KMC berada pada halaman sistem.",
            "post": "Admin KMC mengetahui informasi aduan terbaru beserta rekomendasi kategorinya.",
            "normal": [
                "Admin KMC memilih menu Notifikasi.",
                "Sistem menampilkan daftar notifikasi aduan beserta hasil rekomendasi kategori, subkategori, OPD tujuan, dan prioritas.",
                "Admin menekan salah satu notifikasi untuk melihat detail pesan."
            ],
            "alt": [
                "Tidak ada notifikasi baru di database.",
                "Sistem menampilkan pesan 'Tidak ada notifikasi aduan baru'."
            ]
        },
        {
            "id": "SUC-04",
            "name": "Memverifikasi Kemungkinan Duplikasi",
            "desc": "Mengonfirmasi apakah notifikasi baru merupakan duplikat dari aduan yang sudah ada tiketnya.",
            "actor": "Admin KMC",
            "pre": "Terdapat notifikasi yang ditandai sistem sebagai 'Kemungkinan Duplikasi'.",
            "post": "Notifikasi diarsipkan (jika duplikat) atau dibuatkan tiket (jika bukan duplikat).",
            "normal": [
                "Admin KMC membuka notifikasi yang terdeteksi sebagai kemungkinan duplikasi.",
                "Sistem menampilkan aduan baru bersandingan dengan aduan lama yang mirip.",
                "Admin menekan tombol 'Konfirmasi Duplikat'.",
                "Sistem mengubah status notifikasi menjadi diarsipkan tanpa membuat tiket."
            ],
            "alt": [
                "Admin menilai aduan tersebut bukan duplikat.",
                "Admin menekan tombol 'Bukan Duplikat'.",
                "Sistem membuat tiket baru berdasarkan hasil analisis notifikasi.",
                "Sistem meneruskannya ke OPD terkait."
            ]
        },
        {
            "id": "SUC-05",
            "name": "Mengelola Tiket",
            "desc": "Admin KMC dapat melihat detail tiket, membuat tiket secara manual, atau mengubah atribut tiket seperti kategori, OPD tujuan, dan prioritas.",
            "actor": "Admin KMC",
            "pre": "Admin KMC berada di halaman Manajemen Tiket.",
            "post": "Data tiket berhasil disimpan atau diperbarui di dalam basis data sistem.",
            "normal": [
                "Admin membuka daftar tiket.",
                "Admin memilih tiket yang ingin diubah.",
                "Admin mengubah kategori, subkategori, atau OPD tujuan.",
                "Admin menyimpan perubahan.",
                "Sistem menyimpan pembaruan dan mencatat riwayat perubahan."
            ],
            "alt": [
                "Admin membuat tiket manual dengan menekan tombol 'Buat Tiket'.",
                "Admin mengisi formulir data aduan, pelapor, dan memilih OPD tujuan.",
                "Admin menyimpan data.",
                "Sistem menghasilkan nomor tiket dan menyimpannya."
            ]
        },
        {
            "id": "SUC-06",
            "name": "Mengelola Data dan Akun OPD",
            "desc": "Admin KMC menambah, mengubah, atau menghapus data instansi OPD beserta kredensial akun penggunanya.",
            "actor": "Admin KMC",
            "pre": "Admin KMC berada di halaman Manajemen OPD.",
            "post": "Data dan akun OPD berhasil diperbarui di basis data.",
            "normal": [
                "Admin memilih menu Manajemen OPD.",
                "Sistem menampilkan daftar OPD.",
                "Admin memilih fungsi tambah atau edit data.",
                "Admin memasukkan nama OPD, profil, dan data kredensial akses.",
                "Admin menekan tombol simpan.",
                "Sistem memvalidasi dan menyimpan data."
            ],
            "alt": [
                "Admin memasukkan data yang tidak lengkap atau format email salah.",
                "Sistem menampilkan pesan error validasi.",
                "Admin memperbaiki isian formulir dan menyimpan ulang."
            ]
        },
        {
            "id": "SUC-07",
            "name": "Melihat Statistik Aduan",
            "desc": "Admin KMC melihat laporan atau grafik perkembangan pengelolaan aduan.",
            "actor": "Admin KMC",
            "pre": "Admin KMC telah berhasil masuk ke sistem.",
            "post": "Admin KMC mendapatkan informasi tren pengelolaan tiket.",
            "normal": [
                "Admin memilih menu Statistik/Laporan.",
                "Sistem mengambil data rekapitulasi tiket berdasarkan rentang waktu, status, dan platform sumber.",
                "Sistem menampilkan visualisasi data dalam bentuk grafik atau tabel."
            ],
            "alt": [
                "Tidak ada."
            ]
        },
        {
            "id": "SUC-08",
            "name": "Mengelola Profil",
            "desc": "Pengguna memperbarui data profil pribadi seperti nama, email, foto, dan kata sandi.",
            "actor": "Admin KMC, Pengguna OPD",
            "pre": "Aktor sedang berada di dalam sistem.",
            "post": "Data profil aktor berhasil diperbarui.",
            "normal": [
                "Aktor memilih menu Profil.",
                "Sistem menampilkan data profil aktor saat ini.",
                "Aktor mengubah informasi yang diinginkan atau memasukkan kata sandi baru.",
                "Aktor menyimpan perubahan.",
                "Sistem memvalidasi dan menyimpan pembaruan profil ke database."
            ],
            "alt": [
                "Aktor memasukkan konfirmasi kata sandi baru yang tidak cocok.",
                "Sistem menampilkan peringatan kata sandi tidak cocok.",
                "Aktor mengulang pengisian kata sandi dan menyimpan."
            ]
        },
        {
            "id": "SUC-09",
            "name": "Melihat Dashboard OPD",
            "desc": "Pengguna OPD memantan ringkasan tiket aduan yang menjadi tanggung jawab instansinya.",
            "actor": "Pengguna OPD",
            "pre": "Pengguna OPD telah berhasil masuk ke sistem.",
            "post": "Pengguna OPD melihat ringkasan tiket khusus untuk instansinya.",
            "normal": [
                "Pengguna OPD mengakses halaman utama.",
                "Sistem menyaring data tiket hanya untuk OPD yang bersangkutan.",
                "Sistem menampilkan ringkasan jumlah tiket masuk, diproses, dan selesai."
            ],
            "alt": [
                "Tidak ada."
            ]
        },
        {
            "id": "SUC-10",
            "name": "Melihat Tiket OPD",
            "desc": "Pengguna OPD membaca daftar dan melihat detail aduan yang ditugaskan kepada OPD terkait.",
            "actor": "Pengguna OPD",
            "pre": "Pengguna OPD berada di Dashboard OPD.",
            "post": "Pengguna OPD melihat rincian keluhan, pelapor, dan riwayat penanganan tiket.",
            "normal": [
                "Pengguna OPD memilih menu Daftar Tiket.",
                "Sistem menampilkan seluruh tiket yang ditugaskan ke OPD tersebut.",
                "Pengguna OPD memilih salah satu tiket.",
                "Sistem menampilkan halaman detail tiket berisi deskripsi aduan dan lampiran."
            ],
            "alt": [
                "Pengguna OPD mencoba mengakses tautan URL tiket milik OPD lain.",
                "Sistem menolak akses.",
                "Sistem menampilkan pesan error otoritas bahwa hak akses ditolak."
            ]
        },
        {
            "id": "SUC-11",
            "name": "Memberikan Tanggapan Tiket",
            "desc": "Pengguna OPD mengirimkan balasan, penjelasan, atau solusi terkait aduan ke dalam sistem.",
            "actor": "Pengguna OPD",
            "pre": "Pengguna OPD sedang membuka halaman detail suatu tiket.",
            "post": "Tanggapan tersimpan dalam riwayat tiket dan status diperbarui.",
            "normal": [
                "Pengguna OPD menuliskan pesan tanggapan pada kolom yang tersedia.",
                "Pengguna OPD menambahkan lampiran foto bukti penanganan (jika ada).",
                "Pengguna OPD menekan tombol Kirim Tanggapan.",
                "Sistem menyimpan pesan dan lampiran ke dalam riwayat tiket."
            ],
            "alt": [
                "Pengguna OPD mengunggah file lampiran yang melebihi batas sistem.",
                "Sistem menolak proses dan menampilkan pesan batas maksimal ukuran file.",
                "Pengguna OPD mengganti file dengan ukuran yang sesuai dan mengirim ulang."
            ]
        },
        {
            "id": "SUC-12",
            "name": "Memperbarui Status Tiket",
            "desc": "Pengguna OPD mengubah tahapan progres penanganan tiket (misal: Diproses, Selesai).",
            "actor": "Pengguna OPD",
            "pre": "Pengguna OPD sedang membuka halaman detail tiket yang masih aktif.",
            "post": "Status tiket berubah dan dicatat ke dalam log riwayat penanganan.",
            "normal": [
                "Pengguna OPD memilih opsi status terbaru pada menu pembaruan status.",
                "Pengguna OPD menambahkan catatan pembaruan status.",
                "Pengguna OPD menyimpan pembaruan.",
                "Sistem mengubah status tiket dan menyimpan log riwayat perubahan status."
            ],
            "alt": [
                "Tidak ada."
            ]
        },
        {
            "id": "SUC-13",
            "name": "Melacak Tiket",
            "desc": "Masyarakat melihat perkembangan penanganan aduan mereka melalui portal publik.",
            "actor": "Masyarakat",
            "pre": "Masyarakat berada di portal pelacakan publik sistem.",
            "post": "Sistem menampilkan progres tiket yang dicari.",
            "normal": [
                "Masyarakat memasukkan Nomor Pelacakan (Resi) tiket ke dalam kolom pencarian.",
                "Masyarakat menekan tombol Lacak.",
                "Sistem mencari data tiket berdasarkan nomor pelacakan tersebut di basis data.",
                "Sistem menampilkan status saat ini, OPD yang menangani, dan riwayat singkat tiket kepada masyarakat."
            ],
            "alt": [
                "Masyarakat memasukkan nomor pelacakan yang salah atau tidak terdaftar.",
                "Sistem gagal menemukan data yang cocok.",
                "Sistem menampilkan pesan peringatan bahwa tiket tidak ditemukan."
            ]
        },
        {
            "id": "SUC-14",
            "name": "Melihat Informasi Pemantauan Aduan Publik",
            "desc": "Masyarakat melihat ringkasan data aduan publik yang ditangani secara transparan.",
            "actor": "Masyarakat",
            "pre": "Masyarakat berada di portal publik sistem.",
            "post": "Sistem menampilkan ringkasan statistik pengelolaan aduan secara terbuka.",
            "normal": [
                "Masyarakat mengakses halaman utama pemantauan publik.",
                "Sistem mengambil data ringkasan agregat tiket (jumlah aduan masuk, aduan selesai, dll).",
                "Sistem menampilkan statistik sederhana beserta daftar aduan publik terkini yang disamarkan identitasnya."
            ],
            "alt": [
                "Tidak ada."
            ]
        }
    ]
    
    for idx, sc in enumerate(scenarios, start=1):
        # Create Table
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        
        # Populate headers & metadata
        headers = [
            ("Skenario ID", sc["id"]),
            ("Nama Use Case", sc["name"]),
            ("Deskripsi", sc["desc"]),
            ("Aktor", sc["actor"]),
            ("Pre-kondisi", sc["pre"]),
            ("Post-kondisi", sc["post"])
        ]
        
        for r_idx, (lbl, val) in enumerate(headers):
            row = table.rows[r_idx]
            # Left column
            row.cells[0].paragraphs[0].text = ""
            run_lbl = row.cells[0].paragraphs[0].add_run(lbl)
            run_lbl.bold = True
            # Right column
            row.cells[1].paragraphs[0].text = val
            
        # Skenario Normal
        row_n = table.rows[6]
        row_n.cells[0].paragraphs[0].text = ""
        run_lbl = row_n.cells[0].paragraphs[0].add_run("Skenario Normal")
        run_lbl.bold = True
        
        # Format the list with newline
        normal_txt = ""
        for n_idx, step in enumerate(sc["normal"], start=1):
            normal_txt += f"{n_idx}. {step}"
            if n_idx < len(sc["normal"]):
                normal_txt += "\n"
        row_n.cells[1].paragraphs[0].text = normal_txt
        
        # Skenario Alternatif
        row_a = table.rows[7]
        row_a.cells[0].paragraphs[0].text = ""
        run_lbl = row_a.cells[0].paragraphs[0].add_run("Skenario Alternatif")
        run_lbl.bold = True
        
        alt_txt = ""
        for a_idx, step in enumerate(sc["alt"], start=1):
            if step == "Tidak ada.":
                alt_txt = "-"
            else:
                alt_txt += f"{a_idx}. {step}"
                if a_idx < len(sc["alt"]):
                    alt_txt += "\n"
        row_a.cells[1].paragraphs[0].text = alt_txt
        
        # Add spacing and Table Caption
        caption_p = doc.add_paragraph()
        caption_p.alignment = 1 # Center
        run_cap = caption_p.add_run(f"Tabel 3.{idx}. Skenario Use Case {sc['name']}")
        run_cap.italic = True
        
        # Add space between tables
        doc.add_paragraph()

    doc.save(doc_path)
    print(f"Saved docx to: {doc_path}")

if __name__ == "__main__":
    main()
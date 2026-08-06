from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Cm

OUTPUT = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\Usecase_scenario_fix.docx")

SCENARIOS = [
    ("UC-01", "Login", "Admin KMC, Pengguna OPD", "Memverifikasi identitas pengguna dan memberikan akses sesuai peran.", "Pengguna berada pada halaman login dan memiliki akun terdaftar.", "Pengguna masuk ke dashboard sesuai peran.", "1. Pengguna mengisi username atau email dan kata sandi.\n2. Pengguna menekan tombol Masuk.\n3. Sistem memvalidasi data masuk.\n4. Sistem mengarahkan pengguna ke Dashboard Admin atau Dashboard OPD.", "Jika data masuk tidak sesuai, sistem menampilkan pesan kesalahan dan pengguna tetap berada pada halaman login."),
    ("UC-02", "Melihat Dashboard Admin", "Admin KMC", "Memantau ringkasan aduan, tiket, dan notifikasi.", "Admin KMC telah masuk ke sistem.", "Ringkasan informasi pengelolaan aduan ditampilkan.", "1. Admin membuka Dashboard Admin.\n2. Sistem mengambil ringkasan notifikasi dan tiket.\n3. Sistem menampilkan statistik tiket, notifikasi terbaru, notifikasi prioritas tinggi, tren aduan, dan distribusi sumber aduan.", "Jika belum terdapat data, sistem menampilkan ringkasan dengan nilai kosong."),
    ("UC-03", "Memantau Notifikasi dan Hasil Klasifikasi", "Admin KMC", "Meninjau notifikasi aduan serta rekomendasi hasil pengolahan sistem.", "Admin KMC telah masuk ke sistem.", "Admin mengetahui informasi aduan dan rekomendasi penanganannya.", "1. Admin membuka menu Notifikasi.\n2. Sistem menampilkan daftar notifikasi.\n3. Admin dapat mencari atau menyaring notifikasi.\n4. Admin membuka notifikasi untuk melihat isi aduan serta rekomendasi kategori, subkategori, OPD tujuan, dan prioritas.", "Jika tidak ada notifikasi yang sesuai, sistem menampilkan informasi bahwa data tidak ditemukan."),
    ("UC-04", "Memverifikasi Kemungkinan Duplikasi", "Admin KMC", "Menentukan apakah notifikasi yang ditandai sebagai kemungkinan duplikasi perlu diarsipkan atau dibuatkan tiket.", "Terdapat notifikasi dengan status menunggu verifikasi kemungkinan duplikasi.", "Notifikasi dikonfirmasi sebagai duplikat atau dibuatkan tiket sebagai aduan baru.", "1. Admin membuka notifikasi yang ditandai sebagai kemungkinan duplikasi.\n2. Sistem menampilkan informasi notifikasi dan referensi aduan terkait.\n3. Admin memilih tindakan Bukan Duplikat.\n4. Sistem membuat tiket dari notifikasi tersebut.", "Jika Admin memilih Konfirmasi Duplikat, sistem mengarsipkan notifikasi dan tidak membuat tiket baru."),
    ("UC-05", "Mengelola Tiket", "Admin KMC", "Membuat, melihat, memperbarui, atau menghapus data tiket.", "Admin KMC telah masuk ke sistem.", "Data tiket tersimpan, diperbarui, atau dihapus sesuai tindakan Admin.", "1. Admin membuka menu Tiket.\n2. Sistem menampilkan daftar tiket.\n3. Admin memilih membuat tiket baru atau membuka tiket yang sudah ada.\n4. Admin mengisi atau memperbarui informasi tiket.\n5. Admin menyimpan perubahan.\n6. Sistem memvalidasi dan menyimpan data tiket.", "Jika data wajib belum lengkap atau tidak valid, sistem menampilkan pesan validasi dan meminta Admin memperbaiki data."),
    ("UC-06", "Mengelola Data dan Akun OPD", "Admin KMC", "Mengelola data OPD beserta akun Pengguna OPD.", "Admin KMC telah masuk ke sistem.", "Data OPD dan akun pengguna berhasil ditambahkan, diperbarui, atau dihapus.", "1. Admin membuka menu Manajemen OPD.\n2. Sistem menampilkan daftar OPD dan akun terkait.\n3. Admin memilih tambah, ubah, atau hapus data.\n4. Admin mengisi atau memperbarui data OPD dan akun.\n5. Sistem memvalidasi dan menyimpan perubahan.", "Jika username atau email telah digunakan, sistem menampilkan pesan validasi dan data tidak disimpan."),
    ("UC-07", "Melihat Statistik Aduan", "Admin KMC", "Melihat informasi statistik dan perkembangan pengelolaan aduan.", "Admin KMC telah masuk ke sistem.", "Statistik aduan ditampilkan kepada Admin.", "1. Admin membuka halaman statistik pada Dashboard Admin.\n2. Sistem mengolah ringkasan data aduan dan tiket.\n3. Sistem menampilkan tren, distribusi sumber aduan, dan ringkasan status tiket.", "Jika belum tersedia data, sistem menampilkan grafik atau ringkasan kosong."),
    ("UC-08", "Mengelola Profil Admin", "Admin KMC", "Memperbarui data akun Admin KMC.", "Admin KMC telah masuk ke sistem.", "Data profil Admin berhasil diperbarui.", "1. Admin membuka menu Profil.\n2. Sistem menampilkan data profil.\n3. Admin memperbarui username, email, kata sandi, atau foto profil.\n4. Admin menyimpan perubahan.\n5. Sistem memvalidasi dan menyimpan data.", "Jika username atau email sudah digunakan, sistem menampilkan pesan validasi."),
    ("UC-09", "Melihat Dashboard OPD", "Pengguna OPD", "Memantau ringkasan tiket yang ditugaskan kepada OPD.", "Pengguna OPD telah masuk ke sistem dan terhubung dengan OPD.", "Ringkasan tiket OPD ditampilkan.", "1. Pengguna OPD membuka Dashboard OPD.\n2. Sistem mengambil tiket yang ditugaskan kepada OPD pengguna.\n3. Sistem menampilkan jumlah tiket berdasarkan perkembangan penanganan dan daftar tiket terbaru.", "Jika akun tidak terhubung dengan OPD, sistem menolak akses."),
    ("UC-10", "Melihat Tiket OPD", "Pengguna OPD", "Melihat daftar dan detail tiket yang menjadi tanggung jawab OPD.", "Pengguna OPD telah masuk ke sistem.", "Pengguna OPD melihat tiket yang ditugaskan kepada OPD-nya.", "1. Pengguna OPD membuka menu Tiket.\n2. Sistem menampilkan tiket yang ditugaskan kepada OPD pengguna.\n3. Pengguna dapat mencari atau menyaring tiket.\n4. Pengguna membuka salah satu tiket.\n5. Sistem menampilkan detail tiket dan riwayat penanganan.", "Jika Pengguna OPD mencoba membuka tiket milik OPD lain, sistem menolak akses."),
    ("UC-11", "Memberikan Tanggapan Tiket", "Pengguna OPD", "Mencatat tanggapan OPD terhadap tiket yang ditugaskan.", "Pengguna OPD membuka tiket milik OPD-nya.", "Tanggapan tersimpan pada riwayat tiket.", "1. Pengguna OPD membuka detail tiket.\n2. Pengguna menuliskan tanggapan dan dapat menambahkan lampiran gambar.\n3. Pengguna mengirim tanggapan.\n4. Sistem memvalidasi dan menyimpan tanggapan.\n5. Sistem memperbarui perkembangan tiket.", "Jika tanggapan kosong atau lampiran tidak memenuhi ketentuan, sistem menampilkan pesan validasi."),
    ("UC-12", "Memperbarui Status Tiket", "Pengguna OPD", "Memperbarui perkembangan penanganan tiket.", "Pengguna OPD membuka tiket milik OPD-nya.", "Status dan riwayat tiket diperbarui.", "1. Pengguna OPD membuka detail tiket.\n2. Pengguna memilih status penanganan dan dapat menambahkan catatan atau lampiran.\n3. Pengguna menyimpan pembaruan.\n4. Sistem memvalidasi dan menyimpan perubahan status serta riwayatnya.", "Jika tidak ada perubahan status, catatan, atau lampiran, sistem menampilkan pesan bahwa tidak ada perubahan yang disimpan."),
    ("UC-13", "Mengelola Profil OPD", "Pengguna OPD", "Memperbarui data akun Pengguna OPD.", "Pengguna OPD telah masuk ke sistem.", "Data profil Pengguna OPD berhasil diperbarui.", "1. Pengguna OPD membuka menu Profil.\n2. Sistem menampilkan data profil.\n3. Pengguna memperbarui nama, email, kata sandi, atau foto profil.\n4. Pengguna menyimpan perubahan.\n5. Sistem memvalidasi dan menyimpan data.", "Jika email sudah digunakan atau data tidak valid, sistem menampilkan pesan validasi."),
    ("UC-14", "Melacak Tiket", "Masyarakat", "Melihat perkembangan tiket menggunakan nomor pelacakan tanpa masuk ke sistem.", "Masyarakat berada pada portal publik.", "Informasi tiket yang dicari ditampilkan.", "1. Masyarakat memasukkan nomor pelacakan.\n2. Masyarakat menekan tombol Lacak.\n3. Sistem mencari tiket sesuai nomor pelacakan.\n4. Sistem menampilkan informasi tiket, status, OPD penanganan, riwayat status, dan tanggapan yang tersedia.", "Jika nomor pelacakan tidak ditemukan, sistem menampilkan pesan bahwa tiket tidak ditemukan."),
    ("UC-15", "Melihat Informasi Pemantauan Aduan Publik", "Masyarakat", "Melihat informasi umum pengelolaan aduan pada portal publik.", "Masyarakat mengakses portal publik.", "Informasi pemantauan aduan publik ditampilkan.", "1. Masyarakat membuka portal publik.\n2. Sistem menampilkan ringkasan jumlah laporan, perkembangan penanganan, informasi per OPD, kategori, serta daftar tiket publik.\n3. Masyarakat dapat melakukan pencarian atau penyaringan daftar tiket publik.", "Jika belum terdapat data, sistem menampilkan ringkasan kosong."),
    ("UC-16", "Logout", "Admin KMC, Pengguna OPD", "Mengakhiri sesi pengguna pada sistem.", "Pengguna telah masuk ke sistem.", "Sesi pengguna berakhir dan halaman login ditampilkan.", "1. Pengguna memilih menu Logout.\n2. Sistem mengakhiri sesi pengguna.\n3. Sistem mengarahkan pengguna ke halaman login.", "Tidak ada."),
]


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)


def add_scenario(doc, code, name, actor, purpose, pre, post, main, alt):
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 2"]
    heading.add_run(f"{code} {name}")

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    fields = [
        ("Nama Use Case", name), ("Aktor", actor), ("Tujuan", purpose),
        ("Kondisi Awal", pre), ("Kondisi Akhir", post),
        ("Alur Utama", main), ("Alur Alternatif", alt),
    ]
    for label, value in fields:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, True)
        set_cell_text(cells[1], value)
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("USE CASE SCENARIO\nSISTEM INFORMASI MANAJEMEN ADUAN MULTI CHANNEL KMC")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    doc.add_paragraph()
    intro = doc.add_paragraph(
        "Dokumen ini menjelaskan skenario interaksi pengguna dengan Sistem Informasi Manajemen Aduan Multi Channel KMC. "
        "Skenario disusun berdasarkan tiga aktor pada Use Case Diagram, yaitu Admin KMC, Pengguna OPD, dan Masyarakat. "
        "Proses otomatis seperti pengambilan aduan media sosial, analisis berbantuan AI, pemeriksaan duplikasi, dan pemantauan batas waktu merupakan proses internal sistem sehingga tidak ditetapkan sebagai aktor use case."
    )
    intro.paragraph_format.first_line_indent = Cm(1)
    intro.paragraph_format.line_spacing = 1.5

    doc.add_heading("Daftar Aktor", level=1)
    actors = doc.add_table(rows=1, cols=3)
    actors.style = "Table Grid"
    for cell, label in zip(actors.rows[0].cells, ["No.", "Aktor", "Deskripsi"]):
        set_cell_text(cell, label, True)
    for row in [
        ("1", "Admin KMC", "Pengguna yang mengelola notifikasi, tiket, data dan akun OPD, statistik aduan, serta profil admin."),
        ("2", "Pengguna OPD", "Pengguna dari OPD yang melihat, menanggapi, dan memperbarui status tiket yang ditugaskan kepada OPD-nya."),
        ("3", "Masyarakat", "Pengguna portal publik yang melacak tiket dan melihat informasi pemantauan aduan publik."),
    ]:
        cells = actors.add_row().cells
        for c, value in zip(cells, row): set_cell_text(c, value)

    doc.add_heading("Skenario Use Case", level=1)
    for scenario in SCENARIOS:
        add_scenario(doc, *scenario)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()

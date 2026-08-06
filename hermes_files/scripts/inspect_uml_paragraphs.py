from docx import Document
p = r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_UML.docx"
d = Document(p)
for i, paragraph in enumerate(d.paragraphs):
    text = paragraph.text.strip()
    if 40 <= i <= 66:
        print(f"{i}: {text!r}")
print('tables=', len(d.tables))
for i, table in enumerate(d.tables):
    print('table', i, 'header=', [cell.text.strip() for cell in table.rows[0].cells])

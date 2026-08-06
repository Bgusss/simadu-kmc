from pathlib import Path
from docx import Document

SOURCE = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_ISI.docx")
OUTPUT = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_UML.docx")


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            paragraph._element.remove(run._element)
    else:
        paragraph.add_run(text)


def remove_table(table) -> None:
    table._element.getparent().remove(table._element)


def find_paragraph(document, text: str) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == text:
            return index
    raise ValueError(f"Heading not found: {text}")


def main() -> None:
    document = Document(SOURCE)

    uml_heading_index = find_paragraph(document, "3.2.2 Perancangan Arus Data")
    db_heading_index = find_paragraph(document, "3.2.3 Perancangan Basis Data")

    # Keep numbering intact: 3.2.3 remains the Basis Data section.
    replace_paragraph_text(
        document.paragraphs[uml_heading_index],
        "3.2.2 Perancangan UML (Unified Modeling Language)",
    )

    # Paragraph sequence under the UML heading before the basis-data heading.
    replacements = {
        uml_heading_index + 1: (
            "Perancangan UML (Unified Modeling Language) digunakan untuk "
            "menggambarkan fungsi dan alur utama Sistem Informasi Manajemen Aduan "
            "Multi Channel KMC. Diagram yang digunakan terdiri atas use case diagram "
            "dan activity diagram. Use case diagram menggambarkan interaksi pengguna "
            "dengan sistem, sedangkan activity diagram menggambarkan urutan aktivitas "
            "pada proses pengelolaan aduan dan tindak lanjut tiket."
        ),
        uml_heading_index + 4: (
            "Use case diagram pada Gambar 3.3 menggambarkan fungsi sistem yang dapat "
            "diakses oleh tiga aktor, yaitu Admin KMC, pengguna OPD, dan masyarakat. "
            "Admin KMC bertugas memantau notifikasi dan hasil klasifikasi, memverifikasi "
            "kemungkinan duplikasi, mengelola tiket, mengelola data dan akun OPD, melihat "
            "statistik aduan, serta mengelola profil. Pengguna OPD bertugas melihat tiket "
            "yang ditugaskan kepada OPD-nya, memberikan tanggapan, memperbarui status "
            "tiket, dan mengelola profil. Masyarakat dapat melacak tiket serta melihat "
            "informasi pemantauan aduan yang bersifat publik."
        ),
        uml_heading_index + 5: (
            "Activity diagram pengolahan aduan pada Gambar 3.4 menggambarkan alur aduan "
            "yang diperoleh dari media sosial hingga menjadi tiket. Data aduan yang belum "
            "pernah direkam diperiksa kelayakannya terlebih dahulu. Pesan yang layak "
            "diproses disimpan sebagai notifikasi, kemudian dianalisis untuk memperoleh "
            "rekomendasi penanganan dan diperiksa kemungkinan duplikasinya."
        ),
        uml_heading_index + 8: (
            "Apabila tidak ditemukan kemungkinan duplikasi, sistem membuat tiket, "
            "menetapkan nomor pelacakan, batas waktu penanganan, serta OPD tujuan. "
            "Apabila ditemukan kemungkinan duplikasi, notifikasi menunggu verifikasi "
            "Admin KMC. Admin dapat mengonfirmasi duplikasi sehingga tiket baru tidak "
            "dibuat, atau menyatakan bukan duplikasi sehingga sistem membuat tiket dari "
            "hasil analisis yang tersedia."
        ),
        uml_heading_index + 9: (
            "Activity diagram tindak lanjut tiket dan eskalasi batas waktu penanganan "
            "ditunjukkan pada Gambar 3.5. Diagram ini menggambarkan aktivitas pengguna "
            "OPD dalam melihat, memproses, memberikan tanggapan, dan memperbarui status "
            "tiket yang ditugaskan kepada OPD-nya."
        ),
        uml_heading_index + 12: (
            "Setiap tanggapan dan perubahan status dicatat sebagai riwayat tiket. Sistem "
            "memantau batas waktu penanganan secara berkala. Tiket yang belum memperoleh "
            "respons setelah melewati batas waktu dipindahkan ke tahap disposisi dan diberi "
            "batas waktu penanganan baru. Apabila kembali melewati batas waktu, sistem "
            "mencatat eskalasi dan menaikkan prioritas tiket sampai tiket ditindaklanjuti "
            "atau diselesaikan."
        ),
    }

    for index, text in replacements.items():
        replace_paragraph_text(document.paragraphs[index], text)

    # The actor table was introduced in the previous revision. The two kating examples
    # use a direct narrative + diagrams structure, so remove only this table.
    for table in list(document.tables):
        header = " | ".join(cell.text.strip() for cell in table.rows[0].cells) if table.rows else ""
        if header == "Aktor | Hak Akses Utama":
            remove_table(table)
            break
    else:
        raise ValueError("Actor table not found")

    document.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Paragraphs: {len(document.paragraphs)}; tables: {len(document.tables)}")


if __name__ == "__main__":
    main()

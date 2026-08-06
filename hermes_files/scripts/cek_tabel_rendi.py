from docx import Document
d = Document(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\Laporan Tugas Akhir_Rendiansyah_3042022029.docx")
print("=== DUMP PARAGRAPHS AROUND P992 ===")
for j in range(980, 1000):
    t = d.paragraphs[j].text.strip()
    if t:
        print(f"P{j}: {t}")
print("=== DUMP TABLES AROUND THE CLASS DIAGRAM ===")
# check if there's any table between P980 and P1000
for i, table in enumerate(d.tables):
    # Find which paragraphs are before/after this table to locate it
    # But python-docx doesn't store paragraph indices for tables directly.
    # Let's inspect the text of the first row of all tables.
    print(f"Table {i}: {[cell.text.strip() for cell in table.rows[0].cells[:4]]}")

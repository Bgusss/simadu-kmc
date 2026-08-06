from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_RENDIANSYAH.docx")
OUTPUT = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_ACTIVITY_RENDIANSYAH.docx")


def replace_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            paragraph._element.remove(run._element)
    else:
        paragraph.add_run(text)


def insert_heading_before(paragraph, text: str) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addprevious(element)
    heading = Paragraph(element, paragraph._parent)
    heading.add_run(text).bold = True
    return heading


def main() -> None:
    doc = Document(SOURCE)

    activity_intro = next(
        p for p in doc.paragraphs
        if p.text.strip().startswith("Activity Diagram adalah diagram dalam UML")
    )
    replace_text(
        activity_intro,
        "Activity Diagram adalah diagram dalam UML yang digunakan untuk menggambarkan alur aktivitas atau proses bisnis dalam sistem, mulai dari awal hingga akhir. Diagram ini menunjukkan urutan proses, keputusan, serta kemungkinan percabangan yang terjadi. Pada Sistem Informasi Manajemen Aduan Multi Channel KMC, activity diagram digunakan untuk menggambarkan pengolahan aduan dari media sosial serta tindak lanjut tiket oleh pengguna OPD sampai proses pemantauan batas waktu penanganan.",
    )

    processing_desc = next(
        p for p in doc.paragraphs
        if p.text.strip().startswith("Activity diagram pengolahan aduan pada Gambar 3.4")
    )
    insert_heading_before(processing_desc, "Activity Diagram Pengolahan Aduan dari Media Sosial")
    replace_text(
        processing_desc,
        "Activity Diagram Pengolahan Aduan dari Media Sosial menggambarkan alur aduan sejak data diperoleh dari sumber media sosial hingga terbentuk tiket. Proses dimulai ketika sistem memperoleh data aduan baru dan memeriksa apakah data tersebut telah direkam sebelumnya. Data yang belum pernah direkam kemudian diperiksa kelayakannya. Pesan yang layak diproses disimpan sebagai notifikasi dan dianalisis untuk menghasilkan rekomendasi kategori, subkategori, OPD tujuan, serta prioritas. Selanjutnya, sistem memeriksa kemungkinan duplikasi. Apabila tidak terdapat kemungkinan duplikasi, sistem membuat tiket, menetapkan nomor pelacakan dan batas waktu penanganan, lalu meneruskan tiket kepada OPD tujuan. Apabila terdapat kemungkinan duplikasi, notifikasi menunggu verifikasi Admin KMC. Tampilan alur aktivitas tersebut ditunjukkan pada Gambar 3.4.",
    )

    follow_up_desc = next(
        p for p in doc.paragraphs
        if p.text.strip().startswith("Activity diagram tindak lanjut tiket dan eskalasi")
    )
    insert_heading_before(follow_up_desc, "Activity Diagram Tindak Lanjut Tiket dan Eskalasi SLA")
    replace_text(
        follow_up_desc,
        "Activity Diagram Tindak Lanjut Tiket dan Eskalasi SLA menggambarkan proses penanganan tiket oleh pengguna OPD. Proses dimulai ketika pengguna OPD membuka tiket yang ditugaskan kepada OPD-nya, kemudian membaca detail aduan dan menentukan tindak lanjut. Pengguna OPD dapat memberikan tanggapan, melampirkan bukti penanganan apabila diperlukan, serta memperbarui status tiket. Setiap tanggapan dan perubahan status dicatat sebagai riwayat tiket. Sistem memantau batas waktu penanganan tiket. Apabila tiket belum memperoleh tindak lanjut hingga melewati batas waktu, sistem menjalankan proses disposisi dan eskalasi sesuai ketentuan yang dirancang. Tampilan alur aktivitas tersebut ditunjukkan pada Gambar 3.5.",
    )

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()

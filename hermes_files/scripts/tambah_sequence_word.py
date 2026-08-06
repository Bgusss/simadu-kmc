from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_ACTIVITY_LENGKAP.docx")
OUTPUT = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_UML_RENDIANSYAH_FIX.docx")

def insert_paragraph_after(paragraph, text: str, style=None, bold: bool = False, italic: bool = False) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    new_p = Paragraph(element, paragraph._parent)
    if style:
        new_p.style = style
    run = new_p.add_run(text)
    run.bold = bold
    run.italic = italic
    return new_p

def main() -> None:
    doc = Document(SOURCE)

    # 1. Menemukan letak Sequence Diagram di Word
    # Sequence Diagram dimulai di sub-heading D. Sequence Diagram
    # Di draf saat ini, Gambar 3.10 dan 3.11 adalah Sequence Diagram
    
    # 2. Posisikan dan bersihkan penamaan Sequence Diagram
    # Penomoran yang benar:
    # Gambar 3.10: Sequence Diagram Pengolahan Aduan
    # Gambar 3.11: Sequence Diagram Tindak Lanjut Tiket oleh Pengguna OPD
    # Gambar 3.12: Sequence Diagram Verifikasi Duplikasi oleh Admin KMC
    # Gambar 3.13: Sequence Diagram Eskalasi Prioritas Otomatis
    # Gambar 3.14: Class Diagram Sistem
    # Gambar 3.15: Entity Relationship Diagram (ERD)
    
    # Kita cari Sequence Diagram kedua (Tindak Lanjut)
    t_lanjut_p = next(
        i for i, p in enumerate(doc.paragraphs)
        if "Gambar 3.11 Sequence Diagram Tindak Lanjut Tiket" in p.text
    )
    anchor = doc.paragraphs[t_lanjut_p]
    heading_style = doc.styles["Heading 4"] if "Heading 4" in doc.styles else None

    # Tambahkan Sequence Diagram 3: Verifikasi Duplikasi
    s3_head = insert_paragraph_after(anchor, "Sequence Diagram Verifikasi Duplikasi oleh Admin KMC", style=heading_style, bold=True)
    s3_desc = insert_paragraph_after(
        s3_head,
        "Sequence Diagram Verifikasi Duplikasi oleh Admin KMC menggambarkan alur interaksi ketika Admin KMC melakukan pemeriksaan terhadap notifikasi yang ditandai sebagai duplikat. Admin membuka detail notifikasi, kemudian antarmuka meminta data perbandingan dari pengontrol. Pengontrol memanggil data aduan baru bersandingan dengan aduan pembanding dari basis data untuk ditampilkan di halaman perbandingan. Setelah admin meninjau dan mengirim keputusan (konfirmasi duplikat atau bukan), pengontrol memperbarui status notifikasi di basis data dan secara otomatis membuat tiket baru apabila aduan tersebut terbukti bukan duplikat. Alur interaksi tersebut ditunjukkan pada Gambar 3.12."
    )
    s3_img = insert_paragraph_after(s3_desc, "[ Sisipkan Gambar 3.12 — Sequence Diagram Verifikasi Duplikasi oleh Admin KMC (GAMBAR_3_12_SEQUENCE_VERIFIKASI_DUPLIKASI) ]")
    s3_cap = insert_paragraph_after(s3_img, "Gambar 3.12 Sequence Diagram Verifikasi Duplikasi oleh Admin KMC")

    # Tambahkan Sequence Diagram 4: Eskalasi SLA
    s4_head = insert_paragraph_after(s3_cap, "Sequence Diagram Eskalasi Prioritas Otomatis", style=heading_style, bold=True)
    s4_desc = insert_paragraph_after(
        s4_head,
        "Sequence Diagram Eskalasi Prioritas Otomatis menggambarkan interaksi berulang yang dijalankan oleh sistem untuk memantau batas waktu penanganan tiket. Penjadwal sistem memicu perintah pengecekan secara berkala, lalu command meminta daftar tiket aktif yang belum direspons dari basis data. Command kemudian memproses perhitungan waktu SLA pada model tiket. Apabila waktu penanganan telah melewati batas yang ditentukan, prioritas tiket diperbarui ke tingkat yang lebih tinggi, durasi SLA baru diatur ulang, dan perubahan status tersebut disimpan kembali ke basis data serta dicatat ke dalam log riwayat tiket. Alur interaksi tersebut ditunjukkan pada Gambar 3.13."
    )
    s4_img = insert_paragraph_after(s4_desc, "[ Sisipkan Gambar 3.13 — Sequence Diagram Eskalasi Prioritas Otomatis (GAMBAR_3_13_SEQUENCE_ESKALASI_SLA) ]")
    s4_cap = insert_paragraph_after(s4_img, "Gambar 3.13 Sequence Diagram Eskalasi Prioritas Otomatis")

    # 3. Menyesuaikan Penomoran Gambar di Class Diagram & ERD
    # Karena ada penambahan 2 diagram sequence, Class Diagram bergeser menjadi Gambar 3.14 dan ERD menjadi 3.15
    for p in doc.paragraphs:
        t = p.text
        # Update Class Diagram References
        if "Class diagram pada Gambar 3.12" in t:
            p.text = t.replace("Gambar 3.12", "Gambar 3.14")
        elif "Gambar 3.12 Class Diagram Sistem Informasi" in t:
            p.text = t.replace("Gambar 3.12", "Gambar 3.14")
        elif "[ Sisipkan Gambar 3.12 — Class Diagram" in t:
            p.text = "[ Sisipkan Gambar 3.14 — Class Diagram (GAMBAR_3_14_CLASS_DIAGRAM) ]"
            
        # Update ERD References
        elif "Hubungan antarentitas ditunjukkan pada Gambar 3.13" in t:
            p.text = t.replace("Gambar 3.13", "Gambar 3.15")
        elif "Gambar 3.13 Entity Relationship Diagram" in t:
            p.text = t.replace("Gambar 3.13", "Gambar 3.15")
        elif "[ Sisipkan Gambar 3.13 — Entity Relationship" in t:
            p.text = "[ Sisipkan Gambar 3.15 — Entity Relationship Diagram (GAMBAR_3_15_ERD) ]"

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")

if __name__ == "__main__":
    # Define globally inside script
    OUTPUT_PATH = r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_UML_RENDIANSYAH_FIX.docx"
    main()
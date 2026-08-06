from pathlib import Path
from docx import Document

PATH = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_UML.docx")


def replace(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            paragraph._element.remove(run._element)
    else:
        paragraph.add_run(text)


def remove(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


d = Document(PATH)
start = next(i for i, p in enumerate(d.paragraphs) if p.text.strip() == "3.2.2 Perancangan UML (Unified Modeling Language)")
end = next(i for i, p in enumerate(d.paragraphs) if p.text.strip() == "3.2.3 Perancangan Basis Data")
section = d.paragraphs[start:end]

# Preserve paragraph styles for the first 10 paragraphs in this section.
replacement = [
    "3.2.2 Perancangan UML (Unified Modeling Language)",
    "Perancangan UML (Unified Modeling Language) digunakan untuk menggambarkan fungsi dan alur utama Sistem Informasi Manajemen Aduan Multi Channel KMC. Diagram yang digunakan terdiri atas use case diagram dan activity diagram. Use case diagram menggambarkan interaksi pengguna dengan sistem, sedangkan activity diagram menggambarkan urutan aktivitas pada proses pengelolaan aduan dan tindak lanjut tiket.",
    "[ Sisipkan Gambar 3.3 — Use Case Diagram (GAMBAR_3_3_USE_CASE_DIAGRAM) ]",
    "Gambar 3.3 Use Case Diagram Sistem Informasi Manajemen Aduan Multi Channel KMC",
    "Use case diagram pada Gambar 3.3 menggambarkan fungsi sistem yang dapat diakses oleh tiga aktor, yaitu Admin KMC, pengguna OPD, dan masyarakat. Admin KMC bertugas memantau notifikasi dan hasil klasifikasi, memverifikasi kemungkinan duplikasi, mengelola tiket, mengelola data dan akun OPD, melihat statistik aduan, serta mengelola profil. Pengguna OPD bertugas melihat tiket yang ditugaskan kepada OPD-nya, memberikan tanggapan, memperbarui status tiket, dan mengelola profil. Masyarakat dapat melacak tiket serta melihat informasi pemantauan aduan yang bersifat publik.",
    "[ Sisipkan Gambar 3.4 — Activity Diagram Pengolahan Aduan (GAMBAR_3_4_ACTIVITY_PENGOLAHAN_ADUAN) ]",
    "Gambar 3.4 Activity Diagram Pengolahan Aduan dari Media Sosial",
    "Activity diagram pengolahan aduan pada Gambar 3.4 menggambarkan alur aduan yang diperoleh dari media sosial hingga menjadi tiket. Data aduan yang belum pernah direkam diperiksa kelayakannya terlebih dahulu. Pesan yang layak diproses disimpan sebagai notifikasi, kemudian dianalisis untuk memperoleh rekomendasi penanganan dan diperiksa kemungkinan duplikasinya. Apabila tidak ditemukan kemungkinan duplikasi, sistem membuat tiket, menetapkan nomor pelacakan, batas waktu penanganan, serta OPD tujuan. Apabila ditemukan kemungkinan duplikasi, notifikasi menunggu verifikasi Admin KMC. Admin dapat mengonfirmasi duplikasi sehingga tiket baru tidak dibuat, atau menyatakan bukan duplikasi sehingga sistem membuat tiket dari hasil analisis yang tersedia.",
    "[ Sisipkan Gambar 3.5 — Activity Diagram Tindak Lanjut SLA (GAMBAR_3_5_ACTIVITY_TINDAK_LANJUT_SLA) ]",
    "Gambar 3.5 Activity Diagram Tindak Lanjut Tiket dan Eskalasi SLA",
    "Activity diagram tindak lanjut tiket dan eskalasi batas waktu penanganan pada Gambar 3.5 menggambarkan aktivitas pengguna OPD dalam melihat, memproses, memberikan tanggapan, dan memperbarui status tiket yang ditugaskan kepada OPD-nya. Setiap tanggapan dan perubahan status dicatat sebagai riwayat tiket. Sistem memantau batas waktu penanganan secara berkala. Tiket yang belum memperoleh respons setelah melewati batas waktu dipindahkan ke tahap disposisi dan diberi batas waktu penanganan baru. Apabila kembali melewati batas waktu, sistem mencatat eskalasi dan menaikkan prioritas tiket sampai tiket ditindaklanjuti atau diselesaikan.",
]

if len(section) < len(replacement):
    raise ValueError("Not enough paragraphs in UML section")
for paragraph, text in zip(section, replacement):
    replace(paragraph, text)
for paragraph in section[len(replacement):]:
    remove(paragraph)

d.save(PATH)
print(f"Cleaned: {PATH}")
print("UML paragraphs:", len(replacement))

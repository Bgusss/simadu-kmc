from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_UML_LENGKAP.docx")
OUTPUT = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_RENDIANSYAH.docx")


def replace_text(paragraph, text: str):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            paragraph._element.remove(run._element)
    else:
        paragraph.add_run(text)


def set_heading(paragraph, text: str):
    replace_text(paragraph, text)


def main():
    doc = Document(SOURCE)

    # Shift subsequent headings to make room for 3.2.3 Perancangan UML
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "3.2.2 Perancangan UML (Unified Modeling Language)":
            set_heading(p, "3.2.3 Perancangan UML (Unified Modeling Language)")
        elif t == "3.2.3 Perancangan Basis Data":
            set_heading(p, "3.2.4 Perancangan Basis Data")
        elif t == "3.2.4 Perancangan Antar Muka":
            set_heading(p, "3.2.5 Perancangan Antar Muka")
        elif t == "3.2.5 Perancangan Pengujian Sistem":
            set_heading(p, "3.2.6 Perancangan Pengujian Sistem")
        elif t == "3.2.6 Perancangan Pengujian Penerimaan Pengguna":
            set_heading(p, "3.2.7 Perancangan Pengujian Penerimaan Pengguna")

    # Locate the start of the UML section
    start_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "3.2.3 Perancangan UML (Unified Modeling Language)")

    # 1. Update Intro Paragraph
    intro_p = doc.paragraphs[start_idx + 1]
    replace_text(intro_p, "Dalam perancangan sistem informasi, diperlukan pemodelan untuk menggambarkan alur data dan hubungan antar komponen dalam sistem. Salah satu metode yang digunakan adalah UML (Unified Modeling Language), yaitu standar pemodelan visual yang membantu mendeskripsikan struktur, aktivitas, dan interaksi dalam sistem perangkat lunak. Pada sistem ini, UML digunakan untuk memodelkan interaksi pengguna, alur proses, serta hubungan antar entitas data. Diagram yang digunakan meliputi Use Case Diagram, Use Case Scenario, Activity Diagram, Sequence Diagram, dan Class Diagram.")

    # We will locate and re-arrange paragraphs dynamically.
    # Find the target elements first by scanning doc.paragraphs.
    uc_placeholder = next(p for p in doc.paragraphs if "[ Sisipkan Gambar 3.3" in p.text)
    uc_caption = next(p for p in doc.paragraphs if p.text.startswith("Gambar 3.3 Use Case Diagram"))
    
    sc_desc = next(p for p in doc.paragraphs if "Use case scenario digunakan untuk menjelaskan alur utama" in p.text)
    sc_table_title = next(p for p in doc.paragraphs if p.text.startswith("Tabel 3.2 Use Case Scenario"))
    sc_table = next(t for t in doc.tables if len(t.columns) == 4 and any("Alur Utama" in cell.text for cell in t.rows[0].cells))

    act_placeholder1 = next(p for p in doc.paragraphs if "[ Sisipkan Gambar 3.4" in p.text)
    act_caption1 = next(p for p in doc.paragraphs if p.text.startswith("Gambar 3.4 Activity Diagram"))
    
    act_placeholder2 = next(p for p in doc.paragraphs if "[ Sisipkan Gambar 3.5" in p.text)
    act_caption2 = next(p for p in doc.paragraphs if p.text.startswith("Gambar 3.5 Activity Diagram"))

    seq_placeholder1 = next(p for p in doc.paragraphs if "[ Sisipkan Gambar 3.7" in p.text)
    seq_caption1 = next(p for p in doc.paragraphs if p.text.startswith("Gambar 3.7 Sequence Diagram"))
    
    seq_placeholder2 = next(p for p in doc.paragraphs if "[ Sisipkan Gambar 3.8" in p.text)
    seq_caption2 = next(p for p in doc.paragraphs if p.text.startswith("Gambar 3.8 Sequence Diagram"))

    cls_placeholder = next(p for p in doc.paragraphs if "[ Sisipkan Gambar 3.9" in p.text)
    cls_caption = next(p for p in doc.paragraphs if p.text.startswith("Gambar 3.9 Class Diagram"))

    # We rebuild the structure step-by-step
    current_element = intro_p._p

    def add_p(text, bold_run=None):
        nonlocal current_element
        new_el = OxmlElement("w:p")
        current_element.addnext(new_el)
        new_p = Paragraph(new_el, doc._body)
        new_p.style = doc.styles["Heading 4"] if "Heading 4" in doc.styles else None
        if bold_run:
            new_p.add_run(bold_run).bold = True
            if text:
                new_p.add_run(text)
        else:
            new_p.add_run(text)
        current_element = new_el
        return new_p

    def add_existing_p(p_obj):
        nonlocal current_element
        current_element.addnext(p_obj._element)
        current_element = p_obj._element

    def add_existing_table(t_obj):
        nonlocal current_element
        current_element.addnext(t_obj._element)
        current_element = t_obj._element

    # A. Use Case Diagram
    add_p("", bold_run="A. Use Case Diagram")
    uc_desc = add_p("Use Case Diagram merupakan salah satu jenis diagram dalam Unified Modeling Language (UML) yang digunakan untuk menggambarkan hubungan antara aktor (pengguna) dan fungsi-fungsi utama dalam sistem, serta menunjukkan bagaimana interaksi terjadi antara pengguna dengan fitur-fitur yang tersedia secara menyeluruh. Dalam sistem ini, terdapat tiga aktor utama, yaitu Admin KMC, Pengguna OPD, dan Masyarakat. Admin KMC memiliki akses untuk memantau notifikasi dan hasil klasifikasi, memverifikasi kemungkinan duplikasi, mengelola tiket, mengelola data dan akun OPD, melihat statistik aduan, serta mengelola profil. Pengguna OPD memiliki akses untuk melihat tiket yang ditugaskan kepada OPD-nya, memberikan tanggapan, memperbarui status tiket, dan mengelola profil. Sementara itu, Masyarakat dapat melacak tiket serta melihat informasi pemantauan aduan yang bersifat publik. Dengan adanya Use Case Diagram, perancangan sistem menjadi lebih terarah dan terstruktur karena setiap fitur disesuaikan dengan kebutuhan serta peran pengguna.")
    add_existing_p(uc_placeholder)
    add_existing_p(uc_caption)

    # B. Use Case Scenario
    add_p("", bold_run="B. Use Case Scenario")
    add_existing_p(sc_desc)
    add_existing_p(sc_table_title)
    add_existing_table(sc_table)

    # C. Activity Diagram
    # We must insert a paragraph first to anchor after the table XML element
    p_c_header = add_p("", bold_run="C. Activity Diagram")
    p_c_intro = add_p("Activity Diagram adalah diagram dalam UML yang digunakan untuk menggambarkan alur aktivitas atau proses bisnis dalam sistem, mulai dari awal hingga akhir. Diagram ini menunjukkan urutan proses serta percabangan logika yang terjadi. Pada sistem informasi manajemen aduan, activity diagram menggambarkan alur seperti pengolahan aduan dari media sosial serta proses tindak lanjut tiket dan eskalasi SLA.")
    
    act_desc1 = add_p("Activity diagram pengolahan aduan pada Gambar 3.4 menggambarkan alur aduan yang diperoleh dari media sosial hingga menjadi tiket. Data aduan yang belum pernah direkam diperiksa kelayakannya terlebih dahulu. Pesan yang layak diproses disimpan sebagai notifikasi, dianalisis untuk memperoleh rekomendasi penanganan, dan diperiksa kemungkinan duplikasinya. Apabila tidak ditemukan kemungkinan duplikasi, sistem membuat tiket, menetapkan nomor pelacakan, batas waktu penanganan, serta OPD tujuan. Apabila ditemukan kemungkinan duplikasi, notifikasi menunggu verifikasi Admin KMC.")
    add_existing_p(act_placeholder1)
    add_existing_p(act_caption1)

    act_desc2 = add_p("Activity diagram tindak lanjut tiket dan eskalasi batas waktu penanganan pada Gambar 3.5 menggambarkan aktivitas pengguna OPD dalam melihat, memproses, memberikan tanggapan, dan memperbarui status tiket yang ditugaskan kepada OPD-nya. Setiap tanggapan dan perubahan status dicatat sebagai riwayat tiket. Sistem memantau batas waktu penanganan secara berkala dan meningkatkan prioritas tiket yang belum ditindaklanjuti sesuai ketentuan yang dirancang.")
    add_existing_p(act_placeholder2)
    add_existing_p(act_caption2)

    # D. Sequence Diagram
    add_p("", bold_run="D. Sequence Diagram")
    p_d_intro = add_p("Sequence diagram adalah salah satu jenis diagram UML yang digunakan untuk menggambarkan urutan interaksi antar objek dalam suatu sistem berdasarkan waktu. Diagram ini menunjukkan bagaimana pesan dikirim dari satu objek ke objek lain, dimulai dari atas ke bawah mengikuti alur waktu.")
    
    seq_desc1 = add_p("Sequence diagram pengolahan aduan pada Gambar 3.7 menggambarkan interaksi ketika data aduan diproses menjadi tiket. Sistem menerima data aduan, melakukan penyaringan kelayakan, menyimpan notifikasi yang layak, memperoleh rekomendasi penanganan, serta memeriksa kemungkinan duplikasi. Aduan yang tidak terdeteksi sebagai duplikasi dibuatkan tiket dan diteruskan kepada OPD tujuan.")
    add_existing_p(seq_placeholder1)
    add_existing_p(seq_caption1)

    seq_desc2 = add_p("Sequence diagram tindak lanjut tiket pada Gambar 3.8 menggambarkan interaksi pengguna OPD saat menindaklanjuti tiket. Pengguna OPD membuka tiket yang ditugaskan kepada OPD-nya, melihat detail tiket, lalu memberikan tanggapan atau memperbarui status penanganan. Sistem menyimpan tanggapan dan perubahan status sebagai riwayat tiket.")
    add_existing_p(seq_placeholder2)
    add_existing_p(seq_caption2)

    # E. Class Diagram
    add_p("", bold_run="E. Class Diagram")
    p_e_intro = add_p("Class Diagram merupakan representasi struktur sistem yang menggambarkan bagaimana kelas-kelas dalam sistem saling berhubungan, beserta atribut dan metode yang dimilikinya. Diagram ini digunakan untuk memodelkan struktur internal aplikasi secara statis, serta memberikan pemahaman mengenai desain sistem dari sisi pemrograman berorientasi objek.")

    cls_desc = add_p("Class diagram pada Gambar 3.9 menggambarkan struktur kelas utama yang mendukung sistem. Kelas pengguna terhubung dengan OPD sesuai peran pengguna. Notifikasi berkaitan dengan hasil klasifikasi dan dapat menghasilkan tiket. Tiket terhubung dengan OPD tujuan, tanggapan tiket, serta riwayat perubahan status. Rancangan ini digunakan sebagai acuan struktur objek pada proses pengembangan sistem.")
    add_existing_p(cls_placeholder)
    add_existing_p(cls_caption)

    # Now remove the old un-arranged copies of elements at the bottom to prevent double rendering.
    # Find 3.2.4 Perancangan Basis Data
    db_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "3.2.4 Perancangan Basis Data")
    
    # We will delete all paragraphs between cls_caption's index and db_idx
    # Let's locate cls_caption index in the current list
    cls_cap_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("Gambar 3.9 Class Diagram"))
    
    # Deleting all paragraphs between cls_cap_idx + 1 and db_idx
    to_delete = []
    for idx in range(cls_cap_idx + 1, db_idx):
        to_delete.append(doc.paragraphs[idx])
        
    for p in to_delete:
        p._element.getparent().remove(p._element)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()

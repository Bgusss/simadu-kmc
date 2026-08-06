from docx import Document
d = Document(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\Laporan Tugas Akhir_Rendiansyah_3042022029.docx")
print("=== CARI CONTOH PARAGRAF SEQUENCE LAIN ===")
count = 0
for i in range(890, len(d.paragraphs)):
    t = d.paragraphs[i].text.strip()
    if t.startswith("Sequence Diagram ") and len(t) < 80:
        print(f"P{i}: {t}")
        # Print next 2 paragraphs containing text
        printed = 0
        idx = i + 1
        while printed < 2 and idx < len(d.paragraphs):
            nxt = d.paragraphs[idx].text.strip()
            if nxt:
                print(f"  -> {nxt}")
                printed += 1
            idx += 1
        count += 1
        if count >= 4:
            break

from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_FINAL_UML.docx")
OUTPUT = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_FINAL_SEQUENCE.docx")

def insert_paragraph_after(paragraph, text: str, style=None, bold: bool = False, italic: bool = False) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    new_p = Paragraph(element, paragraph._parent)
    if style:
        new_p.style = style
    run = new_p.add_run(text)
    run.bold = bold
    run.italic = italic
    return new_p

def main() -> None:
    doc = Document(SOURCE)

    # Temukan paragraf akhir dari Sequence Diagram Tindak Lanjut (Gambar 3.11)
    anchor_idx = next(
        i for i, p in enumerate(doc.paragraphs)
        if "Gambar 3.11 Sequence Diagram Tindak Lanjut" in p.text
    )
    anchor = doc.paragraphs[anchor_idx]

    # 1. Sequence Diagram Verifikasi Duplikasi oleh Admin KMC (Gambar 3.12)
    p1_desc = insert_paragraph_after(
        anchor,
        "Sequence diagram verifikasi duplikasi pada Gambar 3.12 menggambarkan alur interaksi ketika Admin KMC melakukan pemeriksaan terhadap notifikasi yang ditandai sebagai kemungkinan duplikasi. Admin KMC membuka halaman detail notifikasi, kemudian sistem mengirim permintaan untuk mengambil data aduan baru bersandingan dengan data aduan lama yang mirip dari basis data. Setelah admin meninjau perbandingan tersebut, admin mengirimkan keputusan verifikasi. Apabila dinyatakan duplikat, sistem memperbarui status notifikasi menjadi diarsipkan tanpa membuat tiket. Apabila dinyatakan bukan duplikat, sistem melanjutkan proses pembuatan tiket dan meneruskannya ke dashboard OPD terkait."
    )
    p1_img = insert_paragraph_after(p1_desc, "[ Sisipkan Gambar 3.12 — Sequence Diagram Verifikasi Duplikasi oleh Admin KMC (GAMBAR_3_12_SEQUENCE_VERIFIKASI_DUPLIKASI) ]")
    p1_cap = insert_paragraph_after(p1_img, "Gambar 3.12 Sequence Diagram Verifikasi Duplikasi oleh Admin KMC")

    # 2. Sequence Diagram Eskalasi Prioritas Otomatis (Gambar 3.13)
    p2_desc = insert_paragraph_after(
        p1_cap,
        "Sequence diagram eskalasi prioritas otomatis pada Gambar 3.13 menggambarkan alur proses latar belakang yang dijalankan oleh penjadwal sistem secara berkala. Penjadwal memicu perintah pengecekan SLA tiket. Sistem kemudian meminta daftar seluruh tiket aktif yang belum selesai dari basis data. Untuk setiap tiket yang telah melewati batas waktu penanganan awal (1x24 jam) tanpa respons OPD, sistem mengubah status menjadi proses disposisi dan memperbarui batas waktu baru. Untuk tiket yang melewati batas waktu berikutnya, sistem menaikkan tingkat prioritas tiket, mencatat log eskalasi, serta memperbarui data batas waktu SLA ke dalam basis data."
    )
    p2_img = insert_paragraph_after(p2_desc, "[ Sisipkan Gambar 3.13 — Sequence Diagram Eskalasi Prioritas Otomatis (GAMBAR_3_13_SEQUENCE_ESKALASI_SLA) ]")
    p2_cap = insert_paragraph_after(p2_img, "Gambar 3.13 Sequence Diagram Eskalasi Prioritas Otomatis")

    # Sesuaikan penomoran Class Diagram (E) dan ERD (Basis Data) setelah penambahan ini.
    # Class Diagram menjadi Gambar 3.14
    # ERD (Basis Data) menjadi Gambar 3.15
    for p in doc.paragraphs[anchor_idx+1:]:
        text = p.text
        if "Class diagram pada Gambar 3.12" in text:
            p.text = text.replace("Gambar 3.12", "Gambar 3.14")
        elif "[ Sisipkan Gambar 3.12 — Class Diagram" in text:
            p.text = text.replace("Gambar 3.12", "Gambar 3.14")
        elif "Gambar 3.12 Class Diagram" in text:
            p.text = text.replace("Gambar 3.12", "Gambar 3.14")
            
        elif "Hubungan antarentitas ditunjukkan pada Gambar 3.13" in text:
            p.text = text.replace("Gambar 3.13", "Gambar 3.15")
        elif "[ Sisipkan Gambar 3.13 — Entity Relationship" in text:
            p.text = text.replace("Gambar 3.13", "Gambar 3.15")
        elif "Gambar 3.13 Entity Relationship Diagram" in text:
            p.text = text.replace("Gambar 3.13", "Gambar 3.15")

    doc.save(OUTPUT_PATH := Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_FINAL_UML.docx"))
    print(f"Saved: {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
from docx import Document
d = Document(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_RENDIANSYAH.docx")
print("=== DUMP ALL FROM P42 TO P78 ===")
for i in range(42, 79):
    t = d.paragraphs[i].text.strip()
    print(f"P{i}: {t}")
print("=== TABLES IN UML SECTION ===")
for i, t in enumerate(d.tables):
    print(f"Table {i}: first cell = {t.cell(0,0).text.strip()}, rows = {len(t.rows)}")

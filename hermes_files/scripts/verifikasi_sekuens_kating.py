from docx import Document
d = Document(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_SEKUENS_KATING.docx")
print("=== VERIFIKASI PERUBAHAN GAYA SEKUENS ===")
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if t.startswith("D. Sequence Diagram"):
        print(f"\n{t}")
        # Cetak 20 paragraf setelah judul
        for j in range(i+1, i+20):
            if j < len(d.paragraphs):
                nxt = d.paragraphs[j].text.strip()
                if nxt:
                    print(f"-> {nxt}")
        break

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_FINAL_UML.docx")
OUTPUT = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_BASIS_DATA.docx")

tables_data = [
    {
        "title": "Tabel 3.x Tabel Database Users",
        "desc": "Tabel ini berisi data akun pengguna sistem, baik Admin KMC maupun pengguna OPD. Setiap akun digunakan untuk otentikasi login dan identifikasi hak akses (role) pada sistem.",
        "headers": ["No", "Atribut", "Tipe Data", "Keterangan"],
        "rows": [
            ["1", "id (PK)", "BigInt", "ID unik pengguna"],
            ["2", "opd_id (FK)", "BigInt", "ID instansi OPD terkait"],
            ["3", "name", "Varchar (255)", "Nama lengkap pengguna"],
            ["4", "username", "Varchar (255)", "Nama pengguna untuk login"],
            ["5", "email", "Varchar (255)", "Alamat surel pengguna"],
            ["6", "password", "Varchar (255)", "Kata sandi terenkripsi"],
            ["7", "role", "Enum", "Hak akses (admin / opd)"],
            ["8", "profile_photo", "Varchar (255)", "Lokasi berkas foto profil"],
            ["9", "created_at", "Timestamp", "Waktu data dibuat"],
            ["10", "updated_at", "Timestamp", "Waktu data diperbarui"]
        ]
    },
    {
        "title": "Tabel 3.x Tabel Database Opds",
        "desc": "Tabel ini menyimpan data referensi master Organisasi Perangkat Daerah (OPD) yang bertanggung jawab dalam penanganan tiket aduan.",
        "headers": ["No", "Atribut", "Tipe Data", "Keterangan"],
        "rows": [
            ["1", "id (PK)", "BigInt", "ID unik OPD"],
            ["2", "name", "Varchar (255)", "Nama instansi OPD"],
            ["3", "created_at", "Timestamp", "Waktu data dibuat"],
            ["4", "updated_at", "Timestamp", "Waktu data diperbarui"]
        ]
    },
    {
        "title": "Tabel 3.x Tabel Database Categories",
        "desc": "Tabel ini berisi referensi kategori utama pengelompokan aduan masyarakat.",
        "headers": ["No", "Atribut", "Tipe Data", "Keterangan"],
        "rows": [
            ["1", "id (PK)", "BigInt", "ID unik kategori"],
            ["2", "name", "Varchar (255)", "Nama kategori aduan"],
            ["3", "created_at", "Timestamp", "Waktu data dibuat"],
            ["4", "updated_at", "Timestamp", "Waktu data diperbarui"]
        ]
    },
    {
        "title": "Tabel 3.x Tabel Database Sub Categories",
        "desc": "Tabel ini menyimpan rincian klasifikasi masalah yang terhubung dengan kategori utama serta OPD yang bertugas menangani masalah tersebut.",
        "headers": ["No", "Atribut", "Tipe Data", "Keterangan"],
        "rows": [
            ["1", "id (PK)", "BigInt", "ID unik subkategori"],
            ["2", "category_id (FK)", "BigInt", "ID kategori utama terkait"],
            ["3", "opd_id (FK)", "BigInt", "ID instansi OPD terkait"],
            ["4", "name", "Varchar (255)", "Nama rincian subkategori"],
            ["5", "created_at", "Timestamp", "Waktu data dibuat"],
            ["6", "updated_at", "Timestamp", "Waktu data diperbarui"]
        ]
    },
    {
        "title": "Tabel 3.x Tabel Database Notifications",
        "desc": "Tabel ini menyimpan data mentah (raw data) pesan masuk dan komentar yang disinkronisasi dari media sosial, sebelum diolah menjadi tiket aduan resmi.",
        "headers": ["No", "Atribut", "Tipe Data", "Keterangan"],
        "rows": [
            ["1", "id (PK)", "BigInt", "ID unik notifikasi"],
            ["2", "duplicate_of_id (FK)", "BigInt", "ID rujukan aduan sebelumnya (jika duplikat)"],
            ["3", "title", "Varchar (255)", "Judul notifikasi (sumber platform)"],
            ["4", "message", "Text", "Isi teks aduan atau komentar"],
            ["5", "sender", "Varchar (255)", "Nama pengirim/akun masyarakat"],
            ["6", "permalink", "Varchar (255)", "Tautan asli ke media sosial"],
            ["7", "is_read", "Boolean", "Status telah dibaca oleh admin"],
            ["8", "duplicate_status", "Varchar (50)", "Status hasil pemeriksaan duplikasi"],
            ["9", "created_at", "Timestamp", "Waktu notifikasi masuk"]
        ]
    },
    {
        "title": "Tabel 3.x Tabel Database AI Classifications",
        "desc": "Tabel ini berfungsi untuk mencatat hasil keluaran kecerdasan buatan (Gemini AI) atas notifikasi masuk, meliputi saran kategori, OPD, hingga tingkat prioritas aduan.",
        "headers": ["No", "Atribut", "Tipe Data", "Keterangan"],
        "rows": [
            ["1", "id (PK)", "BigInt", "ID unik klasifikasi AI"],
            ["2", "notification_id (FK)", "BigInt", "ID notifikasi yang dianalisis"],
            ["3", "suggested_category", "Varchar (255)", "Saran nama kategori"],
            ["4", "suggested_sub_category", "Varchar (255)", "Saran nama subkategori"],
            ["5", "suggested_opds", "Json", "Array daftar saran instansi OPD"],
            ["6", "priority", "Enum", "Tingkat prioritas hasil analisis"],
            ["7", "confidence", "Decimal (5,2)", "Tingkat keyakinan (akurasi) model AI"],
            ["8", "reasoning", "Text", "Penjelasan alasan klasifikasi"],
            ["9", "created_at", "Timestamp", "Waktu data dibuat"]
        ]
    },
    {
        "title": "Tabel 3.x Tabel Database Tickets",
        "desc": "Tabel ini merupakan tabel utama yang menyimpan data tiket penanganan aduan masyarakat yang telah divalidasi dan akan dikelola oleh OPD.",
        "headers": ["No", "Atribut", "Tipe Data", "Keterangan"],
        "rows": [
            ["1", "id (PK)", "BigInt", "ID unik tiket"],
            ["2", "notification_id (FK)", "BigInt", "ID notifikasi asal tiket"],
            ["3", "assigned_opd_id (FK)", "BigInt", "ID instansi OPD penerima tiket"],
            ["4", "ticket_number", "Varchar (255)", "Nomor identifikasi internal tiket"],
            ["5", "tracking_number", "Varchar (255)", "Nomor resi untuk pelacakan publik"],
            ["6", "platform", "Varchar (255)", "Platform sumber asal aduan"],
            ["7", "reporter_name", "Varchar (255)", "Nama pelapor"],
            ["8", "complaint", "Text", "Isi teks keluhan"],
            ["9", "priority", "Enum", "Tingkat prioritas (rendah, sedang, tinggi)"],
            ["10", "status", "Enum", "Status tahapan proses penanganan"],
            ["11", "sla_deadline", "Timestamp", "Batas akhir waktu penanganan tiket"],
            ["12", "created_at", "Timestamp", "Waktu tiket dibuat"]
        ]
    },
    {
        "title": "Tabel 3.x Tabel Database Ticket Responses",
        "desc": "Tabel ini digunakan untuk merekam tanggapan, penjelasan, atau solusi dari pengguna OPD terkait penyelesaian sebuah tiket.",
        "headers": ["No", "Atribut", "Tipe Data", "Keterangan"],
        "rows": [
            ["1", "id (PK)", "BigInt", "ID unik tanggapan"],
            ["2", "ticket_id (FK)", "BigInt", "ID tiket yang ditanggapi"],
            ["3", "user_id (FK)", "BigInt", "ID pengguna yang memberi tanggapan"],
            ["4", "message", "Text", "Isi pesan penjelasan atau solusi"],
            ["5", "attachment", "Varchar (255)", "Lokasi berkas lampiran pendukung"],
            ["6", "created_at", "Timestamp", "Waktu tanggapan diberikan"]
        ]
    },
    {
        "title": "Tabel 3.x Tabel Database Ticket Status Logs",
        "desc": "Tabel ini merekam riwayat seluruh perubahan tahapan (status) tiket yang digunakan untuk pelacakan jejak rekam (audit trail) kinerja penyelesaian aduan.",
        "headers": ["No", "Atribut", "Tipe Data", "Keterangan"],
        "rows": [
            ["1", "id (PK)", "BigInt", "ID unik riwayat status"],
            ["2", "ticket_id (FK)", "BigInt", "ID tiket yang diperbarui"],
            ["3", "changed_by (FK)", "BigInt", "ID pengguna yang memperbarui"],
            ["4", "from_status", "Varchar (255)", "Status tiket sebelum diperbarui"],
            ["5", "to_status", "Varchar (255)", "Status tiket setelah diperbarui"],
            ["6", "note", "Text", "Catatan opsional pembaruan"],
            ["7", "created_at", "Timestamp", "Waktu perubahan status direkam"]
        ]
    }
]

def add_p(element, text, bold_run=None):
    new_el = OxmlElement("w:p")
    element.addnext(new_el)
    new_p = Paragraph(new_el, element.getparent())
    if bold_run:
        new_p.add_run(bold_run).bold = True
        if text:
            new_p.add_run(text)
    else:
        new_p.add_run(text)
    return new_p

def add_table(doc, element, header_data, row_data):
    # Find available table style
    available_styles = [s.name for s in doc.styles if s.type == 3]
    table_style = 'Table Grid' if 'Table Grid' in available_styles else available_styles[0] if available_styles else None
    
    table = doc.add_table(rows=1, cols=len(header_data))
    if table_style:
        table.style = table_style
    
    # Fill headers
    for i, head in enumerate(header_data):
        table.cell(0, i).text = head
        # make it bold
        table.cell(0, i).paragraphs[0].runs[0].bold = True

    # Fill rows
    for row in row_data:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    # Move table to correct location
    element.addnext(table._element)
    return table

def main():
    doc = Document(SOURCE)
    
    # 1. Find anchor paragraph
    anchor_p = None
    for p in doc.paragraphs:
        if "Relasi tersebut dirancang agar data tiket dapat ditelusuri kembali ke notifikasi asal dan hasil analisisnya." in p.text:
            anchor_p = p
            break
            
    if not anchor_p:
        print("Could not find anchor paragraph")
        return
        
    current_element = anchor_p._element

    # 2. Add sub-heading and introduction like Rendiansyah
    # "Struktur Tabel\nPerancangan basis data merupakan tahap penting dalam pengembangan sistem informasi..."
    
    p_header = add_p(current_element, "", bold_run="Struktur Tabel")
    current_element = p_header._element
    
    p_intro = add_p(current_element, "Perancangan basis data merupakan tahap penting dalam pengembangan sistem informasi, karena struktur dan relasi antardata akan memengaruhi efisiensi, kecepatan, dan keakuratan sistem dalam mengelola informasi. Pada sistem ini, perancangan basis data dilakukan untuk memastikan bahwa seluruh data yang dibutuhkan dapat tersimpan dengan baik, mudah diakses, dan dapat saling terhubung antar entitas sesuai kebutuhan sistem. Untuk rancangan kamus data pada sistem ini dapat dilihat pada tabel-tabel berikut.")
    current_element = p_intro._element
    
    # 3. Loop through tables_data to create descriptions and tables
    for i, data in enumerate(tables_data):
        title = data["title"].replace("3.x", f"3.{i+3}") # Table 3.1 is alat & bahan, 3.2 is SUC
        
        # Sub-heading
        p_sub = add_p(current_element, "", bold_run=title.replace(f"Tabel 3.{i+3} ", ""))
        current_element = p_sub._element
        
        # Description
        p_desc = add_p(current_element, data["desc"])
        current_element = p_desc._element
        
        # Table Caption
        p_cap = add_p(current_element, title)
        current_element = p_cap._element
        
        # The Table itself
        tbl = add_table(doc, current_element, data["headers"], data["rows"])
        current_element = tbl._element
        
        # Spacer
        p_space = add_p(current_element, "")
        current_element = p_space._element

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

if __name__ == "__main__":
    main()

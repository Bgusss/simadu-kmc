from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_UML.docx")
OUTPUT = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_UML_LENGKAP.docx")


def replace_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            paragraph._element.remove(run._element)
    else:
        paragraph.add_run(text)


def find_index(doc, text):
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return i
    raise ValueError(f"Tidak ditemukan: {text}")


def insert_after(paragraph, text):
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    new_paragraph.style = paragraph.style
    new_paragraph.add_run(text)
    return new_paragraph


def insert_after_table(table, text, parent):
    new_element = OxmlElement("w:p")
    table._element.addnext(new_element)
    new_paragraph = Paragraph(new_element, parent)
    new_paragraph.add_run(text)
    return new_paragraph


def clone_table_after(doc, table, paragraph):
    cloned = deepcopy(table._element)
    paragraph._p.addnext(cloned)
    return doc.tables[-1]


def set_cell(cell, text):
    replace_text(cell.paragraphs[0], text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def fill_table(table, rows):
    while len(table.rows) < len(rows):
        table.add_row()
    for i, values in enumerate(rows):
        if len(table.rows[i].cells) != len(values):
            raise ValueError("Jumlah kolom tabel tidak sesuai")
        for cell, value in zip(table.rows[i].cells, values):
            set_cell(cell, value)


def main():
    doc = Document(SOURCE)
    start = find_index(doc, "3.2.2 Perancangan UML (Unified Modeling Language)")
    db = find_index(doc, "3.2.3 Perancangan Basis Data")

    # Revisi narasi UML yang sudah ada.
    replacement = {
        start + 1: "Perancangan UML (Unified Modeling Language) digunakan untuk menggambarkan fungsi dan alur utama Sistem Informasi Manajemen Aduan Multi Channel KMC. Diagram yang digunakan meliputi use case diagram, activity diagram, sequence diagram, dan class diagram. Selain itu, use case scenario digunakan untuk menjelaskan alur interaksi utama antara aktor dan sistem.",
        start + 4: "Use case diagram pada Gambar 3.3 menggambarkan fungsi sistem yang dapat diakses oleh tiga aktor, yaitu Admin KMC, pengguna OPD, dan masyarakat. Admin KMC memantau notifikasi dan hasil klasifikasi, memverifikasi kemungkinan duplikasi, mengelola tiket, mengelola data dan akun OPD, melihat statistik aduan, serta mengelola profil. Pengguna OPD melihat tiket yang ditugaskan kepada OPD-nya, memberikan tanggapan, memperbarui status tiket, dan mengelola profil. Masyarakat dapat melacak tiket serta melihat informasi pemantauan aduan yang bersifat publik.",
        start + 7: "Activity diagram pengolahan aduan pada Gambar 3.4 menggambarkan alur aduan yang diperoleh dari media sosial hingga menjadi tiket. Data aduan yang belum pernah direkam diperiksa kelayakannya terlebih dahulu. Pesan yang layak diproses disimpan sebagai notifikasi, dianalisis untuk memperoleh rekomendasi penanganan, dan diperiksa kemungkinan duplikasinya. Apabila tidak ditemukan kemungkinan duplikasi, sistem membuat tiket, menetapkan nomor pelacakan, batas waktu penanganan, serta OPD tujuan. Apabila ditemukan kemungkinan duplikasi, notifikasi menunggu verifikasi Admin KMC.",
        start + 10: "Activity diagram tindak lanjut tiket dan eskalasi batas waktu penanganan pada Gambar 3.5 menggambarkan aktivitas pengguna OPD dalam melihat, memproses, memberikan tanggapan, dan memperbarui status tiket yang ditugaskan kepada OPD-nya. Setiap tanggapan dan perubahan status dicatat sebagai riwayat tiket. Sistem memantau batas waktu penanganan secara berkala dan meningkatkan prioritas tiket yang belum ditindaklanjuti sesuai ketentuan yang dirancang.",
    }
    for index, text in replacement.items():
        replace_text(doc.paragraphs[index], text)

    # Tambahkan Use Case Scenario, Sequence, dan Class Diagram tepat sebelum 3.2.3.
    anchor = doc.paragraphs[db - 1]
    scenario_intro = insert_after(anchor, "Use case scenario digunakan untuk menjelaskan alur utama dari fungsi yang diakses oleh masing-masing aktor. Skenario pada Tabel 3.2 merangkum proses utama pengelolaan aduan, tindak lanjut tiket, dan pelacakan tiket publik.")
    scenario_title = insert_after(scenario_intro, "Tabel 3.2 Use Case Scenario")

    # Buat tabel empat kolom untuk Use Case Scenario dan pindahkan setelah judul tabel.
    scenario_table = doc.add_table(rows=1, cols=4)
    scenario_title._p.addnext(scenario_table._element)
    fill_table(scenario_table, [
        ["No.", "Use Case", "Aktor", "Alur Utama"],
        ["1", "Mengelola Notifikasi dan Tiket", "Admin KMC", "Admin membuka daftar notifikasi, meninjau hasil pengolahan aduan, lalu mengelola atau membuat tiket sesuai kebutuhan."],
        ["2", "Memverifikasi Kemungkinan Duplikasi", "Admin KMC", "Admin meninjau notifikasi yang ditandai sebagai kemungkinan duplikasi dan menentukan apakah notifikasi diarsipkan atau dibuatkan tiket."],
        ["3", "Menindaklanjuti Tiket", "Pengguna OPD", "Pengguna OPD membuka tiket yang ditugaskan kepada OPD-nya, memberikan tanggapan, dan memperbarui status penanganan."],
        ["4", "Melacak Tiket", "Masyarakat", "Masyarakat memasukkan nomor pelacakan untuk melihat informasi perkembangan tiket yang bersifat publik."],
    ])

    # Table is inserted immediately after title. Continue text after its XML element.
    p = insert_after_table(scenario_table, "Sequence diagram pengolahan aduan pada Gambar 3.7 menggambarkan interaksi ketika data aduan diproses menjadi tiket. Sistem menerima data aduan, melakukan penyaringan kelayakan, menyimpan notifikasi yang layak, memperoleh rekomendasi penanganan, serta memeriksa kemungkinan duplikasi. Aduan yang tidak terdeteksi sebagai duplikasi dibuatkan tiket dan diteruskan kepada OPD tujuan.", scenario_title._parent)
    p = insert_after(p, "[ Sisipkan Gambar 3.7 — Sequence Diagram Pengolahan Aduan (GAMBAR_3_7_SEQUENCE_PENGOLAHAN_ADUAN) ]")
    p = insert_after(p, "Gambar 3.7 Sequence Diagram Pengolahan Aduan")
    p = insert_after(p, "Sequence diagram tindak lanjut tiket pada Gambar 3.8 menggambarkan interaksi pengguna OPD saat menindaklanjuti tiket. Pengguna OPD membuka tiket yang ditugaskan kepada OPD-nya, melihat detail tiket, lalu memberikan tanggapan atau memperbarui status penanganan. Sistem menyimpan tanggapan dan perubahan status sebagai riwayat tiket.")
    p = insert_after(p, "[ Sisipkan Gambar 3.8 — Sequence Diagram Tindak Lanjut Tiket OPD (GAMBAR_3_8_SEQUENCE_TINDAK_LANJUT_TIKET) ]")
    p = insert_after(p, "Gambar 3.8 Sequence Diagram Tindak Lanjut Tiket oleh Pengguna OPD")
    p = insert_after(p, "Class diagram pada Gambar 3.9 menggambarkan struktur kelas utama yang mendukung sistem. Kelas pengguna terhubung dengan OPD sesuai peran pengguna. Notifikasi berkaitan dengan hasil klasifikasi dan dapat menghasilkan tiket. Tiket terhubung dengan OPD tujuan, tanggapan tiket, serta riwayat perubahan status. Rancangan ini digunakan sebagai acuan struktur objek pada proses pengembangan sistem.")
    p = insert_after(p, "[ Sisipkan Gambar 3.9 — Class Diagram (GAMBAR_3_9_CLASS_DIAGRAM) ]")
    insert_after(p, "Gambar 3.9 Class Diagram Sistem Informasi Manajemen Aduan Multi Channel KMC")

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Paragraphs: {len(doc.paragraphs)}; tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()

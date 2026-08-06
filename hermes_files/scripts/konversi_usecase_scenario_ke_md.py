from pathlib import Path
from docx import Document

SOURCE = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\Usecase_scenario_fix.docx")
OUTPUT = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\md\Usecase_scenario_fix.md")


def esc(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", "<br>")


def markdown_table(table) -> str:
    rows = [[esc(cell.text.strip()) for cell in row.cells] for row in table.rows]
    header = rows[0]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in rows[1:])
    return "\n".join(lines)


def main():
    doc = Document(SOURCE)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = doc.tables

    lines = [
        "# USE CASE SCENARIO",
        "## Sistem Informasi Manajemen Aduan Multi Channel KMC",
        "",
        paragraphs[2],
        "",
        "## Daftar Aktor",
        "",
        markdown_table(tables[0]),
        "",
        "## Skenario Use Case",
    ]

    scenario_headings = [p for p in paragraphs if p.startswith("UC-")]
    if len(scenario_headings) != 16 or len(tables) != 17:
        raise ValueError(f"Expected 16 scenarios and 17 tables; got {len(scenario_headings)} scenarios and {len(tables)} tables")

    for heading, table in zip(scenario_headings, tables[1:]):
        lines.extend(["", f"### {heading}", "", markdown_table(table)])

    lines.append("")
    OUTPUT.write_text("\n".join(lines), encoding="utf-8")
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()

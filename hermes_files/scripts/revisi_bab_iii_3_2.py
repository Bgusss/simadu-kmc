from copy import deepcopy
from pathlib import Path

from docx import Document

SOURCE = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN.docx")
OUTPUT = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_ISI.docx")


def replace_paragraph_text(paragraph, text: str) -> None:
    """Replace text without changing the paragraph properties or first-run formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            paragraph._element.remove(run._element)
    else:
        paragraph.add_run(text)


def replace_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def set_table_rows(table, rows: list[list[str]]) -> None:
    required = len(rows)
    while len(table.rows) < required:
        table.add_row()

    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        if len(row.cells) != len(values):
            raise ValueError(
                f"Table column mismatch on row {row_index}: "
                f"expected {len(row.cells)}, received {len(values)}"
            )
        for cell, value in zip(row.cells, values):
            replace_cell_text(cell, value)


def main() -> None:
    document = Document(SOURCE)

    expected_headings = {
        34: "3.2 Perancangan Sistem",
        35: "3.2.1 Arsitektur Sistem",
        42: "3.2.2 Perancangan Arus Data",
        57: "3.2.3 Perancangan Basis Data",
        62: "3.2.4 Perancangan Antar Muka",
        68: "3.2.5 Perancangan Pengujian Sistem",
        73: "3.2.6 Perancangan Pengujian Penerimaan Pengguna",
    }
    for index, expected in expected_headings.items():
        actual = document.paragraphs[index].text.strip()
        if actual != expected:
            raise ValueError(f"Unexpected paragraph {index}: {actual!r}")

    replacements = {
        36: (
            "Arsitektur Sistem Informasi Manajemen Aduan Multi Channel KMC "
            "dirancang sebagai sistem informasi berbasis web yang mengintegrasikan "
            "sumber aduan media sosial, pengambilan data, klasifikasi berbantuan AI, "
            "pengelolaan tiket, basis data, dan pengguna sistem. Rancangan arsitektur "
            "ini menggambarkan hubungan antarkomponen secara umum sebagaimana ditunjukkan "
            "pada Gambar 3.2."
        ),
        39: (
            "Aduan bersumber dari mention dan komentar Facebook serta pesan langsung "
            "Instagram. Sistem dirancang untuk melakukan pengambilan data secara berkala "
            "melalui komponen pengambil data. Data yang diperoleh direkam sebagai data "
            "sumber, kemudian diproses dalam aplikasi untuk membentuk notifikasi aduan "
            "dan mendukung pengelolaan tindak lanjut."
        ),
        40: (
            "Setiap pesan yang direkam dirancang melalui proses penyaringan kelayakan "
            "agar pesan yang tidak relevan tidak diteruskan sebagai notifikasi aduan. "
            "Notifikasi yang layak diproses dianalisis oleh layanan AI untuk memberikan "
            "rekomendasi kategori, subkategori, OPD tujuan, dan prioritas penanganan. "
            "Sistem juga memeriksa kemungkinan kesamaan dengan aduan yang telah diproses. "
            "Aduan yang tidak memiliki kemungkinan duplikasi dibuatkan tiket dan "
            "ditetapkan kepada OPD tujuan, sedangkan aduan yang memiliki kemungkinan "
            "duplikasi menunggu verifikasi admin sebelum tiket dibuat."
        ),
        41: (
            "Sistem dirancang untuk diakses oleh tiga jenis pengguna, yaitu Admin KMC, "
            "pengguna OPD, dan masyarakat. Admin KMC dan pengguna OPD mengakses fungsi "
            "sesuai perannya melalui antarmuka web. Masyarakat dapat menggunakan portal "
            "publik untuk melacak perkembangan tiket dan melihat informasi pengelolaan "
            "aduan tanpa perlu masuk ke sistem."
        ),
        43: (
            "Perancangan arus data menggunakan diagram use case dan activity diagram "
            "untuk menggambarkan interaksi aktor serta alur proses utama sistem. Aktor "
            "yang dirancang meliputi Admin KMC, pengguna OPD, masyarakat, sistem scraper, "
            "layanan AI, dan sistem penjadwalan. Rancangan use case sistem ditunjukkan "
            "pada Gambar 3.3."
        ),
        46: (
            "Admin KMC dapat masuk ke sistem, melihat dashboard, memantau notifikasi dan "
            "hasil klasifikasi, memverifikasi kemungkinan duplikasi, mengelola tiket, "
            "mengelola data serta akun OPD, melihat statistik, dan mengelola profil. "
            "Pengguna OPD dapat melihat dashboard dan tiket yang ditugaskan kepada OPD-nya, "
            "memberikan tanggapan, memperbarui status tiket, serta mengelola profil. "
            "Masyarakat dapat melacak dan melihat perkembangan tiket melalui nomor "
            "pelacakan. Sistem scraper, layanan AI, dan sistem penjadwalan mendukung "
            "pengambilan aduan, analisis aduan, serta pemantauan proses secara berkala."
        ),
        49: (
            "Alur pengolahan aduan dari media sosial ditunjukkan pada Gambar 3.4. Proses "
            "dimulai dengan sinkronisasi data aduan dari sumber media sosial. Sistem "
            "merekam data sumber yang belum pernah diperoleh untuk mencegah pengambilan "
            "data yang sama secara berulang. Pesan kemudian diperiksa kelayakannya. Pesan "
            "yang tidak layak diproses tidak diteruskan sebagai notifikasi aduan, sedangkan "
            "pesan yang layak diproses disimpan sebagai notifikasi, dianalisis oleh layanan "
            "AI, dan diperiksa kemungkinan duplikasinya."
        ),
        52: (
            "Apabila tidak ditemukan kemungkinan duplikasi, sistem membuat tiket, "
            "menetapkan nomor pelacakan, batas waktu penanganan, dan OPD tujuan berdasarkan "
            "hasil analisis. Apabila ditemukan kemungkinan duplikasi, sistem menandai "
            "notifikasi untuk diverifikasi oleh Admin KMC. Admin dapat mengonfirmasi bahwa "
            "notifikasi merupakan duplikasi sehingga tiket baru tidak dibuat, atau "
            "menyatakan bukan duplikasi sehingga sistem membuat tiket berdasarkan hasil "
            "analisis yang telah tersedia."
        ),
        56: (
            "Setiap tiket dirancang untuk dikaitkan dengan OPD tujuan. Pengguna OPD dapat "
            "membaca, memproses, memberikan tanggapan, dan menyelesaikan tiket yang "
            "ditugaskan kepada OPD-nya. Setiap tanggapan dan perubahan status dicatat pada "
            "riwayat tiket. Komponen pemantauan memeriksa batas waktu penanganan secara "
            "berkala. Jika tiket yang telah diteruskan kepada OPD belum memperoleh respons "
            "hingga melewati batas waktu, sistem memindahkan tiket ke tahap disposisi dan "
            "menetapkan batas waktu baru. Jika tiket kembali melewati batas waktu tersebut, "
            "sistem mencatat eskalasi dan menaikkan prioritas tiket secara bertahap sampai "
            "tiket ditindaklanjuti atau diselesaikan."
        ),
        58: (
            "Basis data dirancang menggunakan MySQL untuk menyimpan informasi yang "
            "diperlukan sistem secara terstruktur. Entitas utama pada perancangan ini "
            "meliputi pengguna, OPD, kategori, subkategori, notifikasi, hasil klasifikasi "
            "AI, tiket, tanggapan tiket, dan riwayat status tiket. Hubungan antarentitas "
            "ditunjukkan pada Gambar 3.6."
        ),
        61: (
            "Data pengguna membedakan akun Admin KMC dan akun pengguna OPD. Akun pengguna "
            "OPD dihubungkan dengan OPD yang diwakilinya, sedangkan akun admin memiliki "
            "akses pengelolaan sistem. Satu kategori dapat memiliki beberapa subkategori, "
            "dan setiap subkategori dikaitkan dengan OPD tujuan. Setiap notifikasi dapat "
            "memiliki satu hasil klasifikasi AI dan paling banyak satu tiket. Tiket dapat "
            "memiliki banyak tanggapan serta riwayat perubahan status. Relasi tersebut "
            "dirancang agar data tiket dapat ditelusuri kembali ke notifikasi asal dan hasil "
            "analisisnya."
        ),
        63: (
            "Perancangan antar muka bertujuan menyediakan tampilan yang mudah dipahami oleh "
            "setiap pengguna sesuai tugasnya. Antarmuka dirancang menggunakan struktur "
            "tampilan web yang konsisten agar pengguna dapat memperoleh informasi dan "
            "menjalankan fungsi yang diperlukan. Halaman Admin KMC dan pengguna OPD hanya "
            "dapat diakses setelah pengguna masuk menggunakan akun masing-masing, sedangkan "
            "halaman publik dapat diakses langsung untuk pelacakan dan pemantauan tiket."
        ),
        66: (
            "Pada dashboard Admin KMC, notifikasi dengan prioritas tinggi dirancang terpisah "
            "dari notifikasi terbaru agar aduan yang membutuhkan perhatian lebih cepat dapat "
            "terlihat dengan jelas. Rancangan dashboard juga menyediakan pembaruan informasi "
            "secara berkala tanpa pengguna perlu memuat ulang halaman."
        ),
        67: (
            "Pada halaman daftar tiket, setiap tiket dirancang memiliki penanda tingkat "
            "prioritas agar Admin KMC dan pengguna OPD dapat mengenali urgensi penanganan "
            "tanpa harus membuka detail tiket satu per satu."
        ),
        69: (
            "Pengujian sistem dirancang menggunakan metode black box testing. Metode ini "
            "berfokus pada kesesuaian keluaran sistem terhadap masukan dan kebutuhan "
            "fungsional tanpa menilai struktur kode program secara langsung. Rancangan "
            "pengujian mencakup skenario normal, masukan tidak valid, pembatasan hak akses, "
            "serta kondisi khusus pada proses sinkronisasi, analisis aduan, pengelolaan "
            "tiket, dan pemantauan SLA. Pendekatan ini digunakan untuk memastikan setiap "
            "fungsi sistem berjalan sesuai kebutuhan (Mustaqbal et al., 2015)."
        ),
        72: (
            "Hasil pelaksanaan setiap skenario disajikan pada BAB IV. Setiap skenario "
            "dinyatakan berhasil apabila hasil aktual sesuai dengan hasil yang diharapkan "
            "pada rancangan pengujian."
        ),
        74: (
            "Pengujian penerimaan pengguna dirancang untuk mengetahui tingkat penerimaan "
            "dan pemahaman pengguna terhadap aplikasi yang dibangun. Responden yang "
            "direncanakan adalah Admin KMC dan pengguna dari OPD yang menggunakan portal "
            "OPD. Pengujian dilakukan setelah responden mencoba fungsi sesuai perannya, "
            "kemudian mengisi kuesioner."
        ),
        75: (
            "Kuesioner menggunakan skala Likert lima tingkat, yaitu Sangat Tidak Setuju "
            "bernilai 1, Tidak Setuju bernilai 2, Netral bernilai 3, Setuju bernilai 4, "
            "dan Sangat Setuju bernilai 5. Nilai jawaban digunakan untuk menghitung "
            "persentase penerimaan pengguna dengan rumus berikut."
        ),
        79: (
            "Hasil UAT tidak digunakan untuk menilai kepuasan masyarakat terhadap penyelesaian "
            "masalah di lapangan. UAT digunakan untuk mengukur penerimaan pengguna terhadap "
            "aplikasi, kemudahan penggunaannya, serta kesesuaian fungsi aplikasi dengan proses "
            "pengelolaan aduan."
        ),
    }

    for paragraph_index, text in replacements.items():
        replace_paragraph_text(document.paragraphs[paragraph_index], text)

    actor_rows = [
        ["Aktor", "Hak Akses Utama"],
        [
            "Admin KMC",
            "Login, melihat dashboard, memantau notifikasi dan hasil klasifikasi, "
            "memverifikasi duplikasi, mengelola tiket, mengelola data dan akun OPD, "
            "melihat statistik, serta mengelola profil.",
        ],
        [
            "Pengguna OPD",
            "Login, melihat dashboard dan tiket yang ditugaskan kepada OPD-nya, "
            "memberikan tanggapan, memperbarui status tiket, dan mengelola profil.",
        ],
        [
            "Masyarakat",
            "Melacak tiket melalui nomor pelacakan serta melihat informasi perkembangan "
            "dan pemantauan tiket yang bersifat publik.",
        ],
        [
            "Sistem scraper",
            "Mengambil data aduan dari Facebook dan Instagram untuk diteruskan ke proses "
            "pengelolaan aduan dalam sistem.",
        ],
        [
            "Layanan AI",
            "Memberikan rekomendasi klasifikasi, penilaian kelayakan pesan, dan penilaian "
            "kemungkinan duplikasi kepada aplikasi.",
        ],
        [
            "Sistem penjadwalan",
            "Menjalankan proses sinkronisasi data serta pemantauan batas waktu penanganan "
            "secara berkala.",
        ],
    ]
    set_table_rows(document.tables[1], actor_rows)

    interface_rows = [
        ["Halaman", "Pengguna", "Rancangan Fungsi"],
        [
            "Halaman Login",
            "Admin KMC dan Pengguna OPD",
            "Memverifikasi email atau username serta kata sandi, kemudian mengarahkan "
            "pengguna sesuai peran.",
        ],
        [
            "Dashboard Admin",
            "Admin KMC",
            "Menampilkan ringkasan tiket berdasarkan status, notifikasi terbaru, "
            "notifikasi prioritas tinggi, tren aduan, dan distribusi sumber aduan.",
        ],
        [
            "Daftar Notifikasi",
            "Admin KMC",
            "Menampilkan aduan masuk, hasil klasifikasi, status baca, dan kemungkinan "
            "duplikasi. Admin dapat mencari, menyaring, dan memverifikasi notifikasi.",
        ],
        [
            "Manajemen Tiket",
            "Admin KMC",
            "Menampilkan daftar dan detail tiket serta menyediakan fungsi membuat tiket "
            "manual, memperbarui kategori, subkategori, OPD tujuan, dan prioritas tiket.",
        ],
        [
            "Manajemen OPD",
            "Admin KMC",
            "Menambah, mengubah, mencari, dan menghapus data OPD beserta akun pengguna OPD.",
        ],
        [
            "Profil Admin",
            "Admin KMC",
            "Memperbarui username, email, kata sandi, dan foto profil Admin KMC.",
        ],
        [
            "Dashboard OPD",
            "Pengguna OPD",
            "Menampilkan ringkasan dan daftar tiket terbaru yang ditugaskan kepada OPD "
            "pengguna.",
        ],
        [
            "Manajemen Tiket OPD",
            "Pengguna OPD",
            "Menampilkan daftar dan detail tiket OPD, riwayat status, formulir tanggapan, "
            "unggahan lampiran, serta pembaruan status tiket.",
        ],
        [
            "Profil Pengguna OPD",
            "Pengguna OPD",
            "Memperbarui nama, email, kata sandi, dan foto profil pengguna OPD.",
        ],
        [
            "Portal Pelacakan Publik",
            "Masyarakat",
            "Mencari tiket berdasarkan nomor pelacakan, melihat perkembangan tiket, serta "
            "melihat ringkasan pemantauan pengelolaan aduan.",
        ],
    ]
    set_table_rows(document.tables[2], interface_rows)

    black_box_rows = [
        ["No.", "Fitur yang Diuji", "Skenario Pengujian", "Hasil yang Diharapkan"],
        ["1", "Login Admin", "Admin memasukkan kredensial yang benar.", "Sistem mengarahkan Admin KMC ke dashboard admin."],
        ["2", "Login Pengguna OPD", "Pengguna OPD memasukkan kredensial yang benar.", "Sistem mengarahkan pengguna ke dashboard OPD."],
        ["3", "Validasi Login", "Pengguna memasukkan kredensial yang salah atau kosong.", "Sistem menolak login dan menampilkan pesan validasi."],
        ["4", "Otorisasi", "Pengguna OPD membuka halaman khusus Admin KMC.", "Sistem menolak akses ke halaman tersebut."],
        ["5", "Dashboard Sesuai Peran", "Admin KMC atau pengguna OPD berhasil masuk ke sistem.", "Sistem menampilkan dashboard dan data yang sesuai dengan peran pengguna."],
        ["6", "Sinkronisasi Facebook", "Data mention Facebook baru berhasil diperoleh komponen scraper.", "Data sumber direkam satu kali dan aduan layak diproses menjadi notifikasi."],
        ["7", "Sinkronisasi Instagram", "Data pesan Instagram baru berhasil diperoleh komponen scraper.", "Data sumber direkam satu kali dan aduan layak diproses menjadi notifikasi."],
        ["8", "Pencegahan Data Ganda", "Sistem menerima kembali data sumber yang telah direkam.", "Sistem tidak membuat data sumber atau notifikasi ganda."],
        ["9", "Filter Pesan Tidak Layak", "Pesan hanya berisi emoji, reaksi singkat, atau promosi.", "Pesan tidak diteruskan sebagai notifikasi aduan dan tidak dibuatkan tiket."],
        ["10", "Klasifikasi AI", "Aduan valid dikirimkan ke layanan AI.", "Sistem menyimpan rekomendasi kategori, subkategori, OPD, prioritas, tingkat keyakinan, dan alasan."],
        ["11", "Fallback Klasifikasi", "Layanan AI gagal atau menghasilkan respons tidak valid.", "Sistem menggunakan aturan cadangan sehingga proses pengolahan tidak berhenti."],
        ["12", "Deteksi Duplikasi", "Aduan baru memiliki masalah, lokasi, dan objek yang sama dengan aduan bertiket.", "Notifikasi ditandai sebagai kemungkinan duplikasi dan menunggu verifikasi Admin KMC."],
        ["13", "Konfirmasi Duplikasi", "Admin KMC menyatakan notifikasi sebagai duplikasi.", "Notifikasi diarsipkan dan sistem tidak membuat tiket baru."],
        ["14", "Konfirmasi Bukan Duplikasi", "Admin KMC menyatakan notifikasi bukan duplikasi.", "Sistem membuat tiket dari hasil analisis yang telah tersedia."],
        ["15", "Pembuatan Tiket Otomatis", "Aduan valid tidak terdeteksi sebagai duplikasi.", "Sistem membuat nomor pelacakan, tiket, batas waktu penanganan, dan riwayat status awal."],
        ["16", "Pembuatan Tiket Manual", "Admin KMC mengisi data tiket manual yang valid.", "Sistem menyimpan tiket, menetapkan OPD tujuan, dan membuat riwayat status awal."],
        ["17", "Pengelolaan Tiket", "Admin KMC mengubah kategori, subkategori, OPD tujuan, atau prioritas tiket.", "Sistem menyimpan perubahan data tiket dan penugasan OPD."],
        ["18", "Manajemen OPD dan Akun OPD", "Admin KMC menambah, mengubah, atau menghapus data OPD beserta akun pengguna OPD.", "Sistem menyimpan perubahan data sesuai tindakan yang dilakukan."],
        ["19", "Tanggapan OPD", "Pengguna OPD mengirimkan tanggapan yang valid pada tiket miliknya.", "Sistem menyimpan tanggapan dan memperbarui status tiket apabila tiket belum ditutup."],
        ["20", "Pembatasan Tiket OPD", "Pengguna OPD membuka tiket milik OPD lain.", "Sistem menolak akses ke tiket tersebut."],
        ["21", "Profil Admin", "Admin KMC memperbarui data profil dengan data yang valid.", "Sistem menyimpan perubahan profil Admin KMC."],
        ["22", "Profil Pengguna OPD", "Pengguna OPD memperbarui data profil dengan data yang valid.", "Sistem menyimpan perubahan profil pengguna OPD."],
        ["23", "SLA Tahap Pertama", "Tiket yang telah diteruskan kepada OPD belum memperoleh respons hingga melewati batas waktu penanganan.", "Sistem memindahkan tiket ke tahap disposisi dan menetapkan batas waktu baru."],
        ["24", "Eskalasi SLA", "Tiket pada tahap disposisi kembali melewati batas waktu penanganan.", "Sistem mencatat eskalasi dan menaikkan prioritas tiket."],
        ["25", "Pelacakan Publik", "Masyarakat memasukkan nomor pelacakan yang tersedia atau tidak tersedia.", "Sistem menampilkan detail tiket yang sesuai atau pesan bahwa tiket tidak ditemukan."],
        ["26", "Pemantauan Publik", "Masyarakat membuka portal pelacakan tanpa memasukkan nomor pelacakan.", "Sistem menampilkan ringkasan informasi pengelolaan aduan yang dapat diakses publik."],
    ]
    set_table_rows(document.tables[3], black_box_rows)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Paragraphs: {len(document.paragraphs)}")
    print("Tables:", [(len(table.rows), len(table.columns)) for table in document.tables])


if __name__ == "__main__":
    main()

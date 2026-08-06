from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_FINAL_UML.docx")
OUTPUT = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_SEKUENS_KATING.docx")

def replace_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            paragraph._element.remove(run._element)
    else:
        paragraph.add_run(text)

def insert_heading_before(paragraph, text: str) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addprevious(element)
    heading = Paragraph(element, paragraph._parent)
    heading.add_run(text).bold = True
    return heading

def main() -> None:
    doc = Document(SOURCE)

    # 1. Update Sequence Diagram Introduction
    seq_intro = next(
        p for p in doc.paragraphs
        if p.text.strip().startswith("Sequence diagram adalah salah satu jenis diagram UML")
    )
    replace_text(
        seq_intro,
        "Sequence Diagram adalah salah satu jenis diagram UML yang digunakan untuk menggambarkan urutan interaksi antar objek dalam suatu sistem berdasarkan waktu. Diagram ini menunjukkan bagaimana pesan dikirim dari satu objek ke objek lain, dimulai dari atas ke bawah mengikuti alur waktu. Komponen utamanya meliputi aktor, objek/kelas, garis waktu (lifeline), pesan (message), dan balasan pesan (reply message). Pada sistem ini, sequence diagram digunakan untuk menjelaskan alur interaksi pada proses pengolahan aduan, tindak lanjut tiket oleh OPD, verifikasi duplikasi, serta eskalasi prioritas otomatis."
    )

    # 2. Sequence Diagram Pengolahan Aduan (Gambar 3.10)
    seq_desc1 = next(
        p for p in doc.paragraphs
        if p.text.strip().startswith("Sequence diagram pengolahan aduan pada Gambar 3.10")
    )
    insert_heading_before(seq_desc1, "Sequence Diagram Pengolahan Aduan")
    replace_text(
        seq_desc1,
        "Sequence diagram ini menjelaskan alur proses ketika data aduan dari media sosial diproses hingga menjadi tiket. Proses dimulai ketika sistem scraper mengirimkan data aduan baru ke pengontrol sistem. Pengontrol kemudian mengirim permintaan ke layanan AI untuk menyaring kelayakan aduan. Setelah memperoleh hasil bahwa aduan layak diproses, pengontrol menyimpan notifikasi tersebut ke dalam basis data. Selanjutnya, pengontrol meminta layanan AI untuk melakukan klasifikasi kategori, subkategori, OPD tujuan, dan tingkat prioritas. Setelah hasil klasifikasi diterima, pengontrol meminta layanan AI untuk mendeteksi kemungkinan duplikasi dengan aduan lain. Apabila tidak ditemukan duplikasi, pengontrol menyimpan data tiket ke basis data dan tiket diteruskan ke OPD yang ditugaskan. Tampilan sequence diagram dari alur ini dapat dilihat pada Gambar 3.10 berikut."
    )

    # 3. Sequence Diagram Tindak Lanjut Tiket oleh Pengguna OPD (Gambar 3.11)
    seq_desc2 = next(
        p for p in doc.paragraphs
        if p.text.strip().startswith("Sequence diagram tindak lanjut tiket pada Gambar 3.11")
    )
    insert_heading_before(seq_desc2, "Sequence Diagram Tindak Lanjut Tiket oleh Pengguna OPD")
    replace_text(
        seq_desc2,
        "Sequence diagram ini menjelaskan alur proses ketika pengguna OPD memberikan tanggapan atau menyelesaikan aduan yang ditugaskan. Proses diawali saat pengguna OPD membuka tiket melalui antarmuka, kemudian antarmuka mengirim permintaan untuk mengambil data detail tiket ke pengontrol. Pengontrol melakukan pencarian data pada basis data dan mengembalikannya ke antarmuka untuk ditampilkan kepada pengguna. Selanjutnya, pengguna OPD memasukkan pesan tanggapan beserta pembaruan status dan mengirimkannya. Antarmuka meneruskan data tersebut ke pengontrol untuk disimpan ke dalam basis data sebagai log riwayat penanganan tiket. Setelah proses penyimpanan berhasil, sistem menampilkan pesan sukses kepada pengguna OPD. Tampilan sequence diagram dari alur ini dapat dilihat pada Gambar 3.11 berikut."
    )

    # 4. Sequence Diagram Verifikasi Duplikasi oleh Admin KMC (Gambar 3.12)
    seq_desc3 = next(
        p for p in doc.paragraphs
        if p.text.strip().startswith("Sequence diagram verifikasi duplikasi pada Gambar 3.12")
    )
    insert_heading_before(seq_desc3, "Sequence Diagram Verifikasi Duplikasi oleh Admin KMC")
    replace_text(
        seq_desc3,
        "Sequence diagram ini menjelaskan alur proses ketika Admin KMC melakukan pemeriksaan terhadap notifikasi yang ditandai sebagai kemungkinan duplikasi. Proses dimulai saat Admin KMC membuka halaman verifikasi melalui antarmuka. Antarmuka kemudian meminta data perbandingan aduan baru dengan aduan pembanding yang mirip ke pengontrol. Pengontrol mengambil data tersebut dari basis data dan mengembalikannya untuk ditampilkan secara bersandingan. Setelah admin meninjau dan mengirimkan keputusan verifikasi, pengontrol memperbarui status notifikasi. Apabila dikonfirmasi sebagai duplikat, sistem mengarsipkan aduan tanpa membuat tiket. Apabila dikonfirmasi bukan duplikat, pengontrol membuat tiket baru, menyimpannya ke basis data, dan meneruskannya ke dashboard OPD terkait. Tampilan sequence diagram dari alur ini dapat dilihat pada Gambar 3.12 berikut."
    )

    # 5. Sequence Diagram Eskalasi Prioritas Otomatis (Gambar 3.13)
    seq_desc4 = next(
        p for p in doc.paragraphs
        if p.text.strip().startswith("Sequence diagram eskalasi prioritas otomatis pada Gambar 3.13")
    )
    insert_heading_before(seq_desc4, "Sequence Diagram Eskalasi Prioritas Otomatis")
    replace_text(
        seq_desc4,
        "Sequence diagram ini menjelaskan alur proses latar belakang yang dijalankan oleh penjadwal sistem secara berkala untuk memantau SLA tiket. Proses dimulai saat penjadwal sistem memicu pengontrol command untuk melakukan pengecekan berkala. Pengontrol mengirim permintaan ke model tiket untuk mengambil seluruh data tiket aktif dari basis data. Setelah data tiket diterima, pengontrol memeriksa batas waktu penanganan. Apabila tiket telah melewati batas waktu SLA tanpa respons dari OPD, pengontrol menaikkan tingkat prioritas tiket secara otomatis, memperbarui batas waktu penanganan baru, serta menyimpan log riwayat eskalasi status ke dalam basis data. Tampilan sequence diagram dari alur ini dapat dilihat pada Gambar 3.13 berikut."
    )

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

if __name__ == "__main__":
    main()
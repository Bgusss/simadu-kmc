from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SOURCE = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_ACTIVITY_RENDIANSYAH.docx")
OUTPUT = Path(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\BAB_III_METODOLOGI_DAN_PERANCANGAN_REVISI_ACTIVITY_LENGKAP.docx")

def insert_paragraph_after(paragraph, text: str, style=None, bold: bool = False, italic: bool = False) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    new_p = Paragraph(element, paragraph._parent)
    if style:
        new_p.style = style
    run = new_p.add_run(text)
    run.bold = bold
    run.italic = italic
    return new_p

def main() -> None:
    doc = Document(SOURCE)

    # Temukan paragraf akhir dari Activity Diagram SLA (sebelum Sequence Diagram)
    anchor_idx = next(
        i for i, p in enumerate(doc.paragraphs)
        if "Gambar 3.5 Activity Diagram Tindak Lanjut Tiket" in p.text
    )
    anchor = doc.paragraphs[anchor_idx]

    heading_style = doc.styles["Heading 4"] if "Heading 4" in doc.styles else None

    # 1. Activity Diagram Login
    p1_head = insert_paragraph_after(anchor, "Activity Diagram Login dan Logout", style=heading_style, bold=True)
    p1_desc = insert_paragraph_after(
        p1_head,
        "Activity Diagram Login dan Logout menggambarkan alur autentikasi pengguna ke dalam sistem. Proses dimulai saat pengguna mengakses halaman login dan memasukkan kredensial berupa email dan kata sandi. Sistem memvalidasi kredensial tersebut dengan data di basis data. Apabila kredensial tidak valid, sistem mengembalikan pengguna ke halaman login dengan pesan kesalahan. Apabila valid, sistem memeriksa peran (role) pengguna dan mengarahkannya ke dashboard yang sesuai, yaitu Dashboard Admin untuk Admin KMC atau Dashboard OPD untuk Pengguna OPD. Untuk keluar dari sistem, pengguna menekan tombol logout, kemudian sistem mengakhiri sesi dan mengembalikan pengguna ke halaman login. Tampilan alur aktivitas tersebut ditunjukkan pada Gambar 3.6."
    )
    p1_img = insert_paragraph_after(p1_desc, "[ Sisipkan Gambar 3.6 — Activity Diagram Login dan Logout (GAMBAR_3_6_ACTIVITY_LOGIN) ]")
    p1_cap = insert_paragraph_after(p1_img, "Gambar 3.6 Activity Diagram Login dan Logout Sistem")

    # 2. Activity Diagram Pembuatan Tiket Manual
    p2_head = insert_paragraph_after(p1_cap, "Activity Diagram Pembuatan Tiket Manual", style=heading_style, bold=True)
    p2_desc = insert_paragraph_after(
        p2_head,
        "Activity Diagram Pembuatan Tiket Manual menggambarkan alur saat Admin KMC membuat tiket aduan secara langsung tanpa melalui proses sinkronisasi media sosial. Proses dimulai ketika Admin KMC membuka formulir pembuatan tiket, kemudian mengisi data pelapor, isi aduan, kategori, prioritas, dan memilih OPD tujuan. Admin menekan tombol simpan, lalu sistem memvalidasi kelengkapan data. Apabila data tidak valid, sistem meminta admin melengkapi formulir. Apabila valid, sistem menghasilkan nomor pelacakan unik, menetapkan batas waktu penanganan awal (SLA), menyimpan tiket ke basis data, dan secara otomatis meneruskan tiket ke dashboard OPD yang dituju. Tampilan alur aktivitas tersebut ditunjukkan pada Gambar 3.7."
    )
    p2_img = insert_paragraph_after(p2_desc, "[ Sisipkan Gambar 3.7 — Activity Diagram Pembuatan Tiket Manual (GAMBAR_3_7_ACTIVITY_TIKET_MANUAL) ]")
    p2_cap = insert_paragraph_after(p2_img, "Gambar 3.7 Activity Diagram Pembuatan Tiket Manual")

    # 3. Activity Diagram Manajemen OPD
    p3_head = insert_paragraph_after(p2_cap, "Activity Diagram Manajemen Data dan Akun OPD", style=heading_style, bold=True)
    p3_desc = insert_paragraph_after(
        p3_head,
        "Activity Diagram Manajemen Data dan Akun OPD menggambarkan alur proses pengelolaan instansi beserta hak aksesnya oleh Admin KMC. Proses dimulai saat Admin KMC mengakses halaman manajemen OPD. Untuk menambah data, admin mengisi formulir informasi OPD dan kredensial akun, lalu sistem memvalidasi dan menyimpannya. Untuk mengubah atau menghapus data, admin memilih OPD dari daftar, lalu sistem mengeksekusi perubahan atau penghapusan dari basis data setelah mendapat konfirmasi. Proses ini memastikan bahwa setiap OPD memiliki akun yang sah untuk merespons tiket aduan. Tampilan alur aktivitas tersebut ditunjukkan pada Gambar 3.8."
    )
    p3_img = insert_paragraph_after(p3_desc, "[ Sisipkan Gambar 3.8 — Activity Diagram Manajemen OPD (GAMBAR_3_8_ACTIVITY_MANAJEMEN_OPD) ]")
    p3_cap = insert_paragraph_after(p3_img, "Gambar 3.8 Activity Diagram Manajemen Data dan Akun OPD")

    # 4. Activity Diagram Pelacakan Publik
    p4_head = insert_paragraph_after(p3_cap, "Activity Diagram Pelacakan Tiket oleh Masyarakat", style=heading_style, bold=True)
    p4_desc = insert_paragraph_after(
        p4_head,
        "Activity Diagram Pelacakan Tiket oleh Masyarakat menggambarkan alur interaksi publik dalam memantau perkembangan aduan. Proses dimulai ketika masyarakat mengakses portal publik dan memasukkan nomor pelacakan (resi) tiket ke dalam kolom pencarian. Sistem menerima masukan dan mencari data tiket di basis data. Apabila tiket tidak ditemukan, sistem menampilkan pesan bahwa nomor pelacakan tidak valid. Apabila ditemukan, sistem menampilkan detail tiket yang bersifat terbuka, status penanganan saat ini, OPD yang menangani, serta riwayat tanggapan yang telah diberikan. Proses ini mendukung transparansi pengelolaan aduan tanpa memerlukan akses login. Tampilan alur aktivitas tersebut ditunjukkan pada Gambar 3.9."
    )
    p4_img = insert_paragraph_after(p4_desc, "[ Sisipkan Gambar 3.9 — Activity Diagram Pelacakan Publik (GAMBAR_3_9_ACTIVITY_PELACAKAN_PUBLIK) ]")
    p4_cap = insert_paragraph_after(p4_img, "Gambar 3.9 Activity Diagram Pelacakan Tiket oleh Masyarakat")

    # Setelah disisipkan, nomor gambar pada D. Sequence Diagram dan E. Class Diagram harus disesuaikan.
    # Gambar Sequence Diagram yang lama adalah 3.7 dan 3.8, sekarang menjadi 3.10 dan 3.11
    # Gambar Class Diagram yang lama adalah 3.9, sekarang menjadi 3.12
    # ERD yang lama 3.6, sekarang menjadi 3.13
    
    for p in doc.paragraphs[anchor_idx+1:]:
        text = p.text
        if "Sequence diagram pengolahan aduan pada Gambar 3.7" in text:
            p.text = text.replace("Gambar 3.7", "Gambar 3.10")
        elif "[ Sisipkan Gambar 3.7" in text:
            p.text = text.replace("Gambar 3.7", "Gambar 3.10")
        elif "Gambar 3.7 Sequence Diagram Pengolahan Aduan" in text:
            p.text = text.replace("Gambar 3.7", "Gambar 3.10")
            
        elif "Sequence diagram tindak lanjut tiket pada Gambar 3.8" in text:
            p.text = text.replace("Gambar 3.8", "Gambar 3.11")
        elif "[ Sisipkan Gambar 3.8" in text:
            p.text = text.replace("Gambar 3.8", "Gambar 3.11")
        elif "Gambar 3.8 Sequence Diagram Tindak Lanjut Tiket" in text:
            p.text = text.replace("Gambar 3.8", "Gambar 3.11")
            
        elif "Class diagram pada Gambar 3.9" in text:
            p.text = text.replace("Gambar 3.9", "Gambar 3.12")
        elif "[ Sisipkan Gambar 3.9" in text:
            p.text = text.replace("Gambar 3.9", "Gambar 3.12")
        elif "Gambar 3.9 Class Diagram" in text:
            p.text = text.replace("Gambar 3.9", "Gambar 3.12")
            
        elif "Hubungan antarentitas ditunjukkan pada Gambar 3.6" in text:
            p.text = text.replace("Gambar 3.6", "Gambar 3.13")
        elif "[ Sisipkan Gambar 3.6" in text:
            p.text = text.replace("Gambar 3.6", "Gambar 3.13")
        elif "Gambar 3.6 Entity Relationship Diagram" in text:
            p.text = text.replace("Gambar 3.6", "Gambar 3.13")

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

if __name__ == "__main__":
    main()
from docx import Document
d = Document(r"C:\laragon\www\SIMADU-KMC\.hermes\desktop-attachments\Laporan Tugas Akhir_Rendiansyah_3042022029.docx")
print("=== DUMP RENDI CLASS DIAGRAM CONTENT ===")
found = False
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if "Class Diagram" in t and ("Gambar" in t or "Tabel" in t or "4." in t):
        print(f"P{i}: {t}")
        found = True
        # print next 10 non-empty paragraphs
        count = 0
        for j in range(i+1, len(d.paragraphs)):
            text = d.paragraphs[j].text.strip()
            if text:
                print(f"  P{j}: {text}")
                count += 1
                if count >= 10:
                    break
